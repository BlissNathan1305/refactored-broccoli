from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()
doc.add_heading('Discussion of Soil Physicochemical Properties', level=1)

# Define a helper function to add paragraphs with formatting
def add_paragraph(text):
    paragraph = doc.add_paragraph(text)
    run = paragraph.runs[0]
    run.font.size = Pt(11)

# Discussion content
discussion = {
    "Sand (%)": "Sand content varied across locations and depths, indicating differences in soil texture and drainage capacity. Higher sand percentages suggest better aeration and lower water retention, which can influence root development and nutrient availability.",
    "Silt (%)": "Silt levels showed moderate variation, contributing to soil fertility and water retention. Soils with balanced silt content typically support better plant growth due to improved nutrient-holding capacity.",
    "Clay (%)": "Clay content was notably variable, affecting soil structure and permeability. High clay percentages can lead to compacted soils with poor drainage, while low clay levels enhance root penetration.",
    "pH (H₂O)": "Soil pH ranged across locations and depths, reflecting the acidity or alkalinity of the soil. Optimal pH values are crucial for nutrient availability and microbial activity. Deviations from neutral pH may require amendments for optimal crop performance.",
    "E/C (ds/m)": "Electrical conductivity values indicated the salinity levels of the soil. Elevated E/C readings suggest potential salt stress for plants, which can impair water uptake and growth.",
    "OC (%)": "Organic carbon levels varied, representing the amount of decomposed organic matter. Higher OC percentages are beneficial for soil fertility, structure, and microbial life.",
    "OM (%)": "Organic matter content followed similar trends to OC, enhancing soil health and nutrient cycling. Locations with higher OM are likely to support more vigorous plant growth.",
    "T/N (%)": "Total nitrogen percentages were indicative of soil fertility. Nitrogen is essential for vegetative growth, and its availability influences crop yield and quality.",
    "CN:R": "Carbon-to-nitrogen ratios provided insight into the decomposition rate of organic matter. Balanced CN ratios promote efficient nutrient cycling and microbial activity.",
    "Av.P (mg/kg)": "Available phosphorus levels were critical for root development and energy transfer in plants. Variations across depths and locations may necessitate targeted fertilization.",
    "Ca Cmol/kg": "Calcium concentrations influenced soil structure and nutrient uptake. Adequate Ca levels support cell wall development and root health.",
    "Mg Cmol/kg": "Magnesium content was essential for chlorophyll synthesis and enzyme activation. Deficiencies may lead to poor photosynthesis and growth.",
    "K Cmol/kg": "Potassium levels affected water regulation and disease resistance. Soils with sufficient K support better crop resilience and productivity.",
    "Na Cmol/kg": "Sodium concentrations were monitored for potential salinity issues. High Na can disrupt soil structure and nutrient balance.",
    "E/A Cmol/kg": "Exchangeable acidity values reflected the presence of acidic cations. Elevated E/A may require liming to neutralize soil pH.",
    "AL⁺ Cmol/kg": "Aluminum levels were assessed due to their potential toxicity in acidic soils. High AL⁺ can inhibit root growth and nutrient uptake.",
    "H⁺ Cmol/kg": "Hydrogen ion concentrations contributed to soil acidity. Managing H⁺ levels is vital for maintaining a favorable pH range.",
    "TEB Cmol/kg": "Total exchangeable bases indicated the soil's capacity to supply essential cations. Higher TEB values suggest better fertility.",
    "ECEC Cmol/kg": "Effective cation exchange capacity measured the soil's ability to retain and exchange nutrients. High ECEC supports sustained nutrient availability.",
    "BSAT (%)": "Base saturation percentages showed the proportion of exchange sites occupied by basic cations. Higher BSAT values are associated with fertile soils."
}

# Add each discussion point to the document
for parameter, text in discussion.items():
    doc.add_heading(parameter, level=2)
    add_paragraph(text)

# Save the document
doc.save("Soil_Properties_Discussion.docx")
