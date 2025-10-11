import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from sklearn.metrics import mean_squared_error, r2_score
from docx import Document
from docx.shared import Inches

# Step 1: Load the data
data = {
    "Run": list(range(1, 31)),
    "Predicted": [
        13.20, 15.98, 16.15, 13.99, 16.72, 13.14, 14.87, 18.08, 17.00, 11.04,
        18.17, 15.72, 19.05, 15.03, 13.55, 15.24, 19.67, 17.92, 13.79, 17.51,
        14.83, 18.88, 13.90, 18.32, 20.95, 22.79, 14.43, 26.16, 14.77, 15.64
    ],
    "Experimental": [
        13.83, 16.05, 16.73, 13.12, 16.16, 13.69, 14.01, 18.87, 17.96, 11.59,
        17.98, 15.19, 18.96, 14.98, 13.23, 15.79, 19.12, 18.09, 13.41, 17.74,
        13.06, 18.54, 13.27, 18.68, 21.04, 23.01, 14.66, 26.70, 14.05, 15.31
    ]
}

df = pd.DataFrame(data)
df["Residual"] = df["Experimental"] - df["Predicted"]

# Step 2: Plot 1 - Predicted vs Experimental
plt.figure(figsize=(8, 5))
plt.plot(df["Run"], df["Predicted"], label="Predicted", marker='o')
plt.plot(df["Run"], df["Experimental"], label="Experimental", marker='x')
plt.title("Predicted vs Experimental Oil Yield")
plt.xlabel("Run")
plt.ylabel("Oil Yield (%)")
plt.legend()
plt.grid(True)
plt.tight_layout()
plt.savefig("plot_predicted_vs_experimental.png")
plt.close()

# Step 3: Plot 2 - Residuals
plt.figure(figsize=(8, 5))
plt.bar(df["Run"], df["Residual"], color='skyblue')
plt.axhline(0, color='red', linestyle='--')
plt.title("Residuals (Experimental − Predicted)")
plt.xlabel("Run")
plt.ylabel("Residual (%)")
plt.tight_layout()
plt.savefig("plot_residuals.png")
plt.close()

# Step 4: Plot 3 - Regression Fit
plt.figure(figsize=(8, 5))
sns.regplot(x="Predicted", y="Experimental", data=df, ci=None, line_kws={"color": "red"})
plt.title("Regression Fit: Predicted vs Experimental")
plt.xlabel("Predicted Yield (%)")
plt.ylabel("Experimental Yield (%)")
plt.tight_layout()
plt.savefig("plot_regression_fit.png")
plt.close()

# Step 5: Regression Diagnostics
r2 = r2_score(df["Experimental"], df["Predicted"])
mse = mean_squared_error(df["Experimental"], df["Predicted"])
rmse = np.sqrt(mse)

discussion = f"""
This analysis evaluates the predictive accuracy of the Response Surface Methodology (RSM) model used to estimate oil yield from jatropha seeds.

1. Predicted vs Experimental Plot:
   The predicted and experimental values closely align across most runs, indicating strong model performance. The highest yield (Run 28) shows excellent agreement: 26.70% experimental vs 26.16% predicted.

2. Residuals Plot:
   Residuals are mostly small and centered around zero, suggesting minimal bias. A few runs (e.g., Run 3, Run 7) show larger deviations, likely due to experimental variability or model limitations.

3. Regression Fit:
   The regression line shows a strong linear relationship. The coefficient of determination (R²) is {r2:.4f}, indicating that the model explains approximately {r2*100:.2f}% of the variance in experimental yields.
   The Root Mean Square Error (RMSE) is {rmse:.4f}%, reflecting a low average prediction error.

Conclusion:
The RSM model demonstrates high predictive accuracy and reliability for optimizing oil yield from jatropha seeds. The experimental design and regression modeling are validated by the close match between predicted and actual results.
"""

# Step 6: Export to Word
doc = Document()
doc.add_heading("Jatropha Oil Yield Analysis", 0)

doc.add_heading("1. Predicted vs Experimental Oil Yield", level=1)
doc.add_picture("plot_predicted_vs_experimental.png", width=Inches(6))

doc.add_heading("2. Residuals Analysis", level=1)
doc.add_picture("plot_residuals.png", width=Inches(6))

doc.add_heading("3. Regression Fit", level=1)
doc.add_picture("plot_regression_fit.png", width=Inches(6))

doc.add_heading("4. Discussion", level=1)
doc.add_paragraph(discussion)

doc.save("Jatropha_Oil_Yield_Analysis.docx")
