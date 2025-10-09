# Garri Sieve Analysis Report Generator
# Generates a Word (.docx) file summarizing sieve sizes, weights, and percentages

from docx import Document
from docx.shared import Inches

# === Create document ===
doc = Document()
doc.add_heading('Sieve Analysis Report for Garri Milling', level=1)

# === Introduction ===
doc.add_paragraph(
    "This report presents the sieve analysis results for 1000 kg of milled garri, "
    "separated into five particle size categories using standard sieve sizes. "
    "The analysis helps to determine the distribution of coarse, medium, fine, and powdery fractions."
)

# === Table Data ===
data = [
    ["Category", "Sieve Range (mm)", "Description", "Weight Retained (kg)", "% of Total (%)"],
    ["1", "> 1.0", "Coarse granules", "180", "18"],
    ["2", "0.85 – 1.0", "Medium–coarse", "260", "26"],
    ["3", "0.60 – 0.85", "Medium", "340", "34"],
    ["4", "0.30 – 0.60", "Fine", "170", "17"],
    ["5", "< 0.30", "Dust/powder", "50", "5"],
    ["", "", "Total", "1000", "100"]
]

# === Create table ===
doc.add_heading('Table 1: Sieve Analysis Results', level=2)
table = doc.add_table(rows=1, cols=len(data[0]))
table.style = 'Light List Accent 1'

# Add header
hdr_cells = table.rows[0].cells
for i, heading in enumerate(data[0]):
    hdr_cells[i].text = heading

# Add data rows
for row in data[1:]:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = str(val)

# === Observations ===
doc.add_heading('Observations', level=2)
doc.add_paragraph(
    "1. The medium-sized garri fraction (0.60–0.85 mm) is dominant, representing 34% of the total weight, "
    "indicating a well-roasted, commercial-grade product.\n"
    "2. The fine fraction (< 0.30 mm) accounts for 5%, usually classified as powder or dust.\n"
    "3. Coarse particles (> 1.0 mm) form 18% of the total and provide a rough texture preferred in some markets.\n"
    "4. The distribution indicates efficient milling and sieving with minimal losses."
)

# === Save file ===
output_path = "Garri_Sieve_Analysis_Report.docx"
doc.save(output_path)

print(f"Report generated successfully: {output_path}")
