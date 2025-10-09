import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# Data from the PDF
parameters = [
    "pH", "Electrical Conductivity (EC)", "Dissolve Oxygen (DO)", "BOD", "Turbidity", "TSS", "TDS",
    "Phosphate", "Alkalinity", "Salinity", "Chloride", "Acidity", "Nitrate", "COD",
    "Temperature", "Iron (Fe)", "Zinc (Zn)", "Copper (Cu)"
]

units = [
    "", "µ/ds", "mg/L", "mg/L", "NTU", "mg/L", "ppm", "mg/L", "mg/L as CaCo3", "PPT", "mg/L",
    "mg/L", "mg/L", "mg/L", "ºC", "mg/L", "mg/L", "mg/L"
]

locations = ["Nung Uyo", "Ekom Iman", "Afaha Idoro", "Ikot Idaha, Ibiono"]

values = [
    [6.42, 6.37, 6.29, 6.18],
    [16.00, 17.44, 21.41, 19.40],
    [1.004, 1.009, 1.020, 1.040],
    [0.10, 0.10, 0.10, 0.09],
    [0.10, 1.00, 0.30, 1.00],
    [0.0003, 0.0007, 0.0002, 0.0002],
    [0.05, 0.05, 0.05, 0.05],
    [0.002, 0.007, 0.005, 0.003],
    [0.001, 0.001, 0.001, 0.001],
    [0.02, 0.09, 0.04, 0.01],
    [0.211, 0.170, 0.117, 0.092],
    [0.003, 0.003, 0.002, 0.003],
    [0.002, 0.004, 0.002, 0.003],
    [0.009, 0.002, 0.003, 0.004],
    [32, 31, 31, 32],
    [0.070, 0.030, 0.090, 0.010],
    [0.02, 0.02, 0.01, 0.02],
    [0.001, 0.001, 0.008, 0.002]
]

# Create output folder
output_folder = "charts"
os.makedirs(output_folder, exist_ok=True)

# Create Word document
doc = Document()
doc.add_heading("Water Quality Analysis Across Four Locations", 0)

# Generate charts and discussion
for i, param in enumerate(parameters):
    plt.figure(figsize=(6, 4))
    plt.bar(locations, values[i], color='skyblue')
    plt.title(f"{param} Levels")
    plt.ylabel(f"Concentration ({units[i]})" if units[i] else "Concentration")
    plt.tight_layout()

    chart_path = os.path.join(output_folder, f"{param.replace(' ', '_')}.png")
    plt.savefig(chart_path)
    plt.close()

    # Add chart to doc
    doc.add_heading(param, level=1)
    doc.add_picture(chart_path, width=Inches(5.5))

    # Add discussion
    doc.add_paragraph(f"**Discussion:** The concentration of {param} varies across the four sampled locations. "
                      f"Nung Uyo recorded a value of {values[i][0]}{units[i]}, Ekom Iman had {values[i][1]}{units[i]}, "
                      f"Afaha Idoro showed {values[i][2]}{units[i]}, and Ikot Idaha, Ibiono had {values[i][3]}{units[i]}. "
                      f"This variation may be attributed to differences in local environmental conditions, anthropogenic activities, "
                      f"and natural water sources. Further investigation could help determine the specific causes of these differences.")

# Save document
doc.save("Water_Quality_Analysis.docx")
print("Document exported successfully as 'Water_Quality_Analysis.docx'")
