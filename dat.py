import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D
from sklearn.preprocessing import PolynomialFeatures
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score, mean_squared_error
from scipy import stats
from scipy.optimize import minimize
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import warnings
warnings.filterwarnings('ignore')

# Data from Table 4.1
data = {
    'Run': list(range(1, 31)),
    'X1': [0,-1,1,-1,1,-1,1,-1,1,0,0,-2,0,0,0,0,0,0,2,0,0,-1,1,-1,1,-1,1,-1,1,0],
    'X2': [0,-1,-1,1,1,-1,-1,1,1,0,-2,0,0,0,0,0,0,0,0,2,0,-1,-1,1,1,-1,-1,1,1,0],
    'X3': [0,-1,-1,-1,-1,1,1,1,1,-2,0,0,0,0,0,0,0,0,0,0,2,-1,-1,-1,-1,1,1,1,1,0],
    'X4': [-2,-1,-1,-1,-1,-1,-1,-1,-1,0,0,0,0,0,0,0,0,0,0,0,0,1,1,1,1,1,1,1,1,2],
    'MC': [10,8,12,8,12,8,12,8,12,6,14,6,14,8,10,6,12,14,6,12,8,14,6,4,10,14,6,12,12,14],
    'Ht': [25,20,35,25,20,20,15,20,20,15,20,15,35,30,25,35,25,15,20,30,20,15,15,35,20,15,25,30,35,30],
    'HT': [70,60,60,60,60,80,80,80,80,50,70,70,90,70,70,90,80,70,70,70,90,60,50,70,90,80,60,90,80,70],
    'SET': [60,90,90,60,90,60,90,150,120,60,150,90,120,90,60,90,180,150,60,90,60,180,60,90,120,180,90,180,60,90],
    'Oil_Yield': [13.83,16.05,16.73,13.12,16.16,13.69,14.01,18.87,17.96,11.59,17.98,15.19,18.96,14.98,13.23,15.79,19.12,18.09,13.41,17.74,13.06,18.54,13.27,18.68,21.04,23.01,14.66,26.70,14.05,15.31]
}

df = pd.DataFrame(data)

# Prepare features (coded values)
X_coded = df[['X1', 'X2', 'X3', 'X4']].values
y = df['Oil_Yield'].values

# Create polynomial features (second order)
poly = PolynomialFeatures(degree=2, include_bias=True)
X_poly = poly.fit_transform(X_coded)

# Fit the model
model = LinearRegression()
model.fit(X_poly, y)

# Predictions
y_pred = model.predict(X_poly)

# Model statistics
r2 = r2_score(y, y_pred)
adj_r2 = 1 - (1-r2)*(len(y)-1)/(len(y)-X_poly.shape[1]-1)
mse = mean_squared_error(y, y_pred)
rmse = np.sqrt(mse)

# Get feature names
feature_names = poly.get_feature_names_out(['X1', 'X2', 'X3', 'X4'])

# Coefficients
coefficients = pd.DataFrame({
    'Term': feature_names,
    'Coefficient': model.coef_
})

# ANOVA
ss_total = np.sum((y - np.mean(y))**2)
ss_regression = np.sum((y_pred - np.mean(y))**2)
ss_residual = np.sum((y - y_pred)**2)

df_regression = X_poly.shape[1] - 1
df_residual = len(y) - X_poly.shape[1]
df_total = len(y) - 1

ms_regression = ss_regression / df_regression
ms_residual = ss_residual / df_residual

f_statistic = ms_regression / ms_residual
p_value = 1 - stats.f.cdf(f_statistic, df_regression, df_residual)

# Create Word document
doc = Document()
doc.add_heading('Jatropha Oil Extraction Analysis Report', 0)

# Model Summary
doc.add_heading('1. Regression Model Summary', 1)
doc.add_paragraph(f'R-squared: {r2:.4f}')
doc.add_paragraph(f'Adjusted R-squared: {adj_r2:.4f}')
doc.add_paragraph(f'RMSE: {rmse:.4f}')
doc.add_paragraph(f'MSE: {mse:.4f}')

# Coefficients
doc.add_heading('2. Model Coefficients', 1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Term'
hdr_cells[1].text = 'Coefficient'

for idx, row in coefficients.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Term'])
    row_cells[1].text = f"{row['Coefficient']:.6f}"

# ANOVA Table
doc.add_heading('3. Analysis of Variance (ANOVA)', 1)
anova_table = doc.add_table(rows=4, cols=6)
anova_table.style = 'Light Grid Accent 1'

# Headers
headers = ['Source', 'SS', 'DF', 'MS', 'F-value', 'P-value']
for i, header in enumerate(headers):
    anova_table.rows[0].cells[i].text = header

# Regression row
anova_table.rows[1].cells[0].text = 'Regression'
anova_table.rows[1].cells[1].text = f'{ss_regression:.4f}'
anova_table.rows[1].cells[2].text = f'{df_regression}'
anova_table.rows[1].cells[3].text = f'{ms_regression:.4f}'
anova_table.rows[1].cells[4].text = f'{f_statistic:.4f}'
anova_table.rows[1].cells[5].text = f'{p_value:.6f}'

# Residual row
anova_table.rows[2].cells[0].text = 'Residual'
anova_table.rows[2].cells[1].text = f'{ss_residual:.4f}'
anova_table.rows[2].cells[2].text = f'{df_residual}'
anova_table.rows[2].cells[3].text = f'{ms_residual:.4f}'

# Total row
anova_table.rows[3].cells[0].text = 'Total'
anova_table.rows[3].cells[1].text = f'{ss_total:.4f}'
anova_table.rows[3].cells[2].text = f'{df_total}'

doc.add_page_break()

# Generate contour and surface plots
doc.add_heading('4. Response Surface Plots', 1)

# Define actual value ranges
ranges = {
    'X1': (6, 14),   # MC
    'X2': (50, 90),  # HT
    'X3': (15, 35),  # Ht
    'X4': (60, 180)  # SET
}

labels = {
    'X1': 'Moisture Content (%)',
    'X2': 'Heating Temperature (°C)',
    'X3': 'Heating Time (min)',
    'X4': 'Extraction Time (min)'
}

# Convert actual to coded values
def actual_to_coded(actual, var):
    if var == 'X1':  # MC
        return (actual - 10) / 2
    elif var == 'X2':  # HT
        return (actual - 70) / 10
    elif var == 'X3':  # Ht
        return (actual - 25) / 5
    elif var == 'X4':  # SET
        return (actual - 120) / 30

# Interaction pairs to plot
plot_pairs = [
    ('X1', 'X2', 0, 0),  # MC vs HT, fixing Ht=25, SET=120
    ('X1', 'X4', 0, 0),  # MC vs SET, fixing HT=70, Ht=25
    ('X2', 'X4', 0, 0),  # HT vs SET, fixing MC=10, Ht=25
    ('X3', 'X4', 0, 0),  # Ht vs SET, fixing MC=10, HT=70
]

plot_count = 1
for var1, var2, fix1, fix2 in plot_pairs:
    # Create mesh grid in actual values
    v1_actual = np.linspace(ranges[var1][0], ranges[var1][1], 50)
    v2_actual = np.linspace(ranges[var2][0], ranges[var2][1], 50)
    V1, V2 = np.meshgrid(v1_actual, v2_actual)
    
    # Convert to coded values
    V1_coded = actual_to_coded(V1, var1)
    V2_coded = actual_to_coded(V2, var2)
    
    # Create prediction grid
    Z = np.zeros_like(V1)
    
    for i in range(V1.shape[0]):
        for j in range(V1.shape[1]):
            # Set up input vector
            input_coded = [0, 0, 0, 0]
            
            if var1 == 'X1':
                input_coded[0] = V1_coded[i, j]
            elif var1 == 'X2':
                input_coded[1] = V1_coded[i, j]
            elif var1 == 'X3':
                input_coded[2] = V1_coded[i, j]
            elif var1 == 'X4':
                input_coded[3] = V1_coded[i, j]
                
            if var2 == 'X1':
                input_coded[0] = V2_coded[i, j]
            elif var2 == 'X2':
                input_coded[1] = V2_coded[i, j]
            elif var2 == 'X3':
                input_coded[2] = V2_coded[i, j]
            elif var2 == 'X4':
                input_coded[3] = V2_coded[i, j]
            
            # Transform and predict
            input_poly = poly.transform([input_coded])
            Z[i, j] = model.predict(input_poly)[0]
    
    # Contour plot
    fig, ax = plt.subplots(figsize=(10, 7))
    contour = ax.contourf(V1, V2, Z, levels=15, cmap='viridis')
    plt.colorbar(contour, ax=ax, label='Oil Yield (%)')
    ax.contour(V1, V2, Z, levels=10, colors='white', linewidths=0.5, alpha=0.4)
    ax.set_xlabel(labels[var1], fontsize=12, fontweight='bold')
    ax.set_ylabel(labels[var2], fontsize=12, fontweight='bold')
    ax.set_title(f'Contour Plot: {labels[var1]} vs {labels[var2]}', fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.3)
    
    # Save to buffer
    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    
    # Add to document
    doc.add_heading(f'4.{plot_count} {labels[var1]} vs {labels[var2]}', 2)
    doc.add_picture(buf, width=Inches(6))
    
    # 3D Surface plot
    fig = plt.figure(figsize=(12, 8))
    ax = fig.add_subplot(111, projection='3d')
    surf = ax.plot_surface(V1, V2, Z, cmap='viridis', alpha=0.9, edgecolor='none')
    ax.set_xlabel(labels[var1], fontsize=11, fontweight='bold')
    ax.set_ylabel(labels[var2], fontsize=11, fontweight='bold')
    ax.set_zlabel('Oil Yield (%)', fontsize=11, fontweight='bold')
    ax.set_title(f'3D Surface: {labels[var1]} vs {labels[var2]}', fontsize=14, fontweight='bold')
    fig.colorbar(surf, ax=ax, shrink=0.5, aspect=5, label='Oil Yield (%)')
    ax.view_init(elev=25, azim=45)
    
    # Save to buffer
    buf = BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    
    # Add to document
    doc.add_picture(buf, width=Inches(6))
    doc.add_page_break()
    
    plot_count += 1

# Optimization
doc.add_heading('5. Optimization Results', 1)

def objective(x_coded):
    x_poly = poly.transform([x_coded])
    return -model.predict(x_poly)[0]  # Negative for maximization

# Bounds (coded values: -2 to 2)
bounds = [(-2, 2), (-2, 2), (-2, 2), (-2, 2)]

# Initial guess
x0 = [0, 0, 0, 0]

# Optimize
result = minimize(objective, x0, method='SLSQP', bounds=bounds)

optimal_coded = result.x
optimal_yield = -result.fun

# Convert to actual values
optimal_mc = 10 + 2 * optimal_coded[0]
optimal_ht = 70 + 10 * optimal_coded[1]
optimal_heating_time = 25 + 5 * optimal_coded[2]
optimal_set = 120 + 30 * optimal_coded[3]

doc.add_paragraph('Optimal Process Parameters (from optimization):')
doc.add_paragraph(f'  • Moisture Content: {optimal_mc:.2f} %')
doc.add_paragraph(f'  • Heating Temperature: {optimal_ht:.2f} °C')
doc.add_paragraph(f'  • Heating Time: {optimal_heating_time:.2f} min')
doc.add_paragraph(f'  • Extraction Time: {optimal_set:.2f} min')
doc.add_paragraph(f'  • Predicted Maximum Oil Yield: {optimal_yield:.2f} %')

doc.add_paragraph('')
doc.add_paragraph('Comparison with Experimental Maximum:')
doc.add_paragraph(f'  • Experimental maximum oil yield: 26.70% (Run 28)')
doc.add_paragraph(f'  • Experimental conditions: MC=12%, HT=90°C, Ht=30min, SET=180min')
doc.add_paragraph(f'  • Optimization prediction: {optimal_yield:.2f}%')
doc.add_paragraph(f'  • Difference: {abs(optimal_yield - 26.70):.2f}%')

# Model equation
doc.add_page_break()
doc.add_heading('6. Regression Model Equation', 1)
doc.add_paragraph('Oil Yield (%) = ')
equation_text = f'{model.coef_[0]:.4f}'
for i, name in enumerate(feature_names[1:], 1):
    coef = model.coef_[i]
    sign = '+' if coef >= 0 else '-'
    equation_text += f' {sign} {abs(coef):.4f}*{name}'

# Add equation in parts (Word has character limits)
for i in range(0, len(equation_text), 100):
    doc.add_paragraph(equation_text[i:i+100])

# Save document
doc.save('Jatropha_Oil_Extraction_Analysis.docx')
print("Analysis complete! Document saved as 'Jatropha_Oil_Extraction_Analysis.docx'")
print(f"\nOptimal conditions found:")
print(f"Moisture Content: {optimal_mc:.2f} %")
print(f"Heating Temperature: {optimal_ht:.2f} °C")
print(f"Heating Time: {optimal_heating_time:.2f} min")
print(f"Extraction Time: {optimal_set:.2f} min")
print(f"Predicted Maximum Oil Yield: {optimal_yield:.2f} %")
