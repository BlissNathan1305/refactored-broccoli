"""
Chapter Five: Conclusion and Recommendations
Exports: Chapter5_Conclusion_and_Recommendations.docx
Dependencies:
    pip install python-docx
"""

from docx import Document

# Create document
doc = Document()
doc.add_heading("CHAPTER FIVE", level=1)
doc.add_heading("CONCLUSION AND RECOMMENDATIONS", level=2)

# --- Conclusion ---
doc.add_heading("5.1 Conclusion", level=3)
doc.add_paragraph(
    "This study evaluated the sieve characteristics of milled garri and milled millet samples, "
    "each analyzed in three replications to determine the distribution of fine, very fine, and powdery fractions. "
    "The findings revealed that both products were dominated by the fine fraction (0.45–0.60 mm), representing more than half "
    "of the total mass in each case — approximately 58% for garri and 57% for millet. The very fine fraction (0.30–0.45 mm) "
    "accounted for about 33%, while the powdery fraction (<0.30 mm) contributed below 10% of the total weight."
)

doc.add_paragraph(
    "The ANOVA results showed statistically significant differences (p < 0.01) among the size fractions, confirming that "
    "milling and sieving effectively separate particles into distinct categories. The low coefficients of variation (below 5%) "
    "across replications indicate high reproducibility and consistency in the milling process."
)

doc.add_paragraph(
    "The findings demonstrate that both garri and millet milling operations were highly efficient, producing uniform and smooth "
    "textures desirable for consumer use and industrial applications. The predominance of fine granules suggests optimized equipment "
    "performance and uniform processing parameters, making the products suitable for direct consumption, packaging, or further value-added processing."
)

# --- Recommendations ---
doc.add_heading("5.2 Recommendations", level=3)

recommendations = [
    "Optimization of Milling Equipment: Regular calibration of milling and sieving machines is recommended to maintain uniform particle sizes and minimize powder loss.",
    "Process Standardization: Establishing a controlled milling time and speed can improve consistency across batches and reduce the proportion of powdery residue.",
    "Quality Control and Grading: Implement routine sieve analysis as a quality control step in garri and millet processing plants to ensure uniformity of texture and compliance with market preferences.",
    "Adoption of Advanced Sieving Technology: Using mechanical or vibratory sieves with mesh sizes between 0.30 mm and 0.60 mm can improve separation precision and product classification.",
    "Further Research: Future work should investigate the influence of roasting temperature, moisture content, and particle size on consumer acceptability, reconstitution properties, and shelf stability.",
    "Training and Awareness: Small-scale processors should receive training on the importance of sieve uniformity, as it affects product quality, market price, and processing efficiency."
]

for rec in recommendations:
    doc.add_paragraph(rec, style="List Number")

# --- Overall Implication ---
doc.add_heading("5.3 Overall Implication", level=3)
doc.add_paragraph(
    "This study provides a data-driven framework for improving post-milling quality control in the processing of garri and millet. "
    "By understanding particle distribution and statistical variability, producers can achieve more consistent textures, better yield efficiency, "
    "and enhanced product quality — contributing significantly to food processing standards and economic sustainability in local agro-industrial production."
)

# Save document
output_path = "Chapter5_Conclusion_and_Recommendations.docx"
doc.save(output_path)

print(f"✅ Chapter Five successfully exported as: {output_path}")
