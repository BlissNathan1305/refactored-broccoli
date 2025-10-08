import pandas as pd
import numpy as np
import statsmodels.formula.api as smf
from scipy.optimize import minimize
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D
from docx import Document
from docx.shared import Inches
import os # To handle temporary image files

# --- 1. Data Extraction and Preparation ---
# Data extracted from Source 50: Table 4.1
data = {
    'Run': list(range(1, 31)),
    'X1': [0, -1, 1, -1, 1, -1, 1, -1, 1, 0, 0, -2, 0, 0, 0, 0, 0, 0, 2, 0, 0, -1, 1, -1, 1, -1, 1, -1, 1, 0],
    'X2': [0, -1, -1, 1, 1, -1, -1, 1, 1, 0, -2, 0, 0, 0, 0, 0, 0, 0, 0, 2, 0, -1, -1, 1, 1, -1, -1, 1, 1, 0],
    'X3': [0, -1, -1, -1, -1, 1, 1, 1, 1, -2, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 2, -1, -1, -1, -1, 1, 1, 1, 1, 0],
    'X4': [-2, -1, -1, -1, -1, -1, -1, -1, -1, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 1, 1, 1, 1, 1, 1, 1, 1, 2],
    'MC (%)': [10, 8, 12, 8, 12, 8, 12, 8, 12, 6, 14, 6, 14, 8, 10, 6, 12, 14, 6, 12, 8, 14, 6, 4, 10, 14, 6, 12, 12, 14],
    'Ht (mins)': [25, 20, 35, 25, 20, 20, 15, 20, 20, 15, 20, 15, 35, 30, 25, 35, 25, 15, 20, 30, 20, 15, 15, 35, 20, 15, 25, 30, 35, 30],
    'HT (°C)': [70, 60, 60, 60, 60, 80, 80, 80, 80, 50, 70, 70, 90, 70, 70, 90, 80, 70, 70, 70, 90, 60, 50, 70, 90, 80, 60, 90, 80, 70],
    'SET (mins)': [60, 90, 90, 60, 90, 60, 90, 150, 120, 120, 150, 90, 120, 90, 60, 90, 180, 150, 60, 90, 60, 180, 60, 90, 120, 180, 90, 180, 60, 90],
    'Y_Exp': [13.83, 16.05, 16.73, 13.12, 16.16, 13.69, 14.01, 18.87, 17.96, 11.59, 17.98, 15.19, 18.96, 14.98, 13.23, 15.79, 19.12, 18.09, 13.41, 17.74, 13.06, 18.54, 13.27, 18.68, 21.04, 23.01, 14.66, 26.70, 14.05, 15.31]
}
df = pd.DataFrame(data)

# --- 2. Model Fitting (Second-Order Polynomial) ---
# X1=Moisture Content, X2=Heating Temperature, X3=Heating Time, X4=Extraction Time

# Full quadratic formula for 4 factors using coded values
rsm_formula = (
    'Y_Exp ~ X1 + X2 + X3 + X4 + ' +
    'I(X1**2) + I(X2**2) + I(X3**2) + I(X4**2) + ' +
    'X1:X2 + X1:X3 + X1:X4 + X2:X3 + X2:X4 + X3:X4'
)

# Fit the model
model = smf.ols(formula=rsm_formula, data=df)
results = model.fit()

# Extract coefficients for the optimization function
coefficients = results.params

# Function to predict yield using the fitted model (coded values)
def predict_yield_coded(X1, X2, X3, X4, coeffs):
    # Intercept (B0)
    Y_pred = coeffs['Intercept']

    # Linear terms (Bi*Xi)
    Y_pred += coeffs['X1'] * X1
    Y_pred += coeffs['X2'] * X2
    Y_pred += coeffs['X3'] * X3
    Y_pred += coeffs['X4'] * X4

    # Quadratic terms (Bii*Xi^2)
    Y_pred += coeffs['I(X1 ** 2)'] * (X1**2)
    Y_pred += coeffs['I(X2 ** 2)'] * (X2**2)
    Y_pred += coeffs['I(X3 ** 2)'] * (X3**2)
    Y_pred += coeffs['I(X4 ** 2)'] * (X4**2)

    # Interaction terms (Bij*Xi*Xj)
    Y_pred += coeffs['X1:X2'] * X1 * X2
    Y_pred += coeffs['X1:X3'] * X1 * X3
    Y_pred += coeffs['X1:X4'] * X1 * X4
    Y_pred += coeffs['X2:X3'] * X2 * X3
    Y_pred += coeffs['X2:X4'] * X2 * X4
    Y_pred += coeffs['X3:X4'] * X3 * X4

    return Y_pred

# --- 3. Optimization to Maximize Predicted Oil Yield ---

# Function to minimize (negative of the predicted oil yield)
def predicted_yield_neg(X):
    X1, X2, X3, X4 = X
    return -predict_yield_coded(X1, X2, X3, X4, coefficients)

# Constraints: -2 <= Xi <= 2 for all factors
bounds = [(-2, 2), (-2, 2), (-2, 2), (-2, 2)]

# Initial guess (start at the best experimental point, Run 28: X = [-1, 1, 1, 1] for 26.70%)
initial_guess = [-1, 1, 1, 1]

# Optimization using the L-BFGS-B method for bounded problems
optimization_result = minimize(
    predicted_yield_neg,
    initial_guess,
    method='L-BFGS-B',
    bounds=bounds
)

# Optimized Coded Factors
X_opt = optimization_result.x
Y_opt = -optimization_result.fun

# Convert coded values to actual values (from Source 16)
MC_opt = 10 + 2 * X_opt[0]
HTemp_opt = 70 + 10 * X_opt[1]
HTime_opt = 25 + 5 * X_opt[2]
SET_opt = 120 + 30 * X_opt[3]

# Prepare results for the DOCX function
opt_results_dict = {
    'Y_opt': Y_opt,
    'MC_opt': MC_opt,
    'HTemp_opt': HTemp_opt,
    'HTime_opt': HTime_opt,
    'SET_opt': SET_opt
}

# --- 4. Visualization Functions ---

# Helper function to get actual factor labels
def get_actual_label(code):
    if code == 'X1': return 'Moisture Content (%wb)'
    if code == 'X2': return 'Heating Temperature (°C)'
    if code == 'X3': return 'Heating Time (min)'
    if code == 'X4': return 'Extraction Time (min)'
    return code

def generate_rsm_plots(X_var_code, Y_var_code, fixed_var_codes, coeffs):
    """
    Generates a 2D contour plot and a 3D surface plot for two varying factors,
    while fixing the other two at their center point (0).
    Saves plots as temporary image files.
    """
    
    # Define the range for the varying factors
    x_range = np.linspace(-2, 2, 50)
    y_range = np.linspace(-2, 2, 50)
    X, Y = np.meshgrid(x_range, y_range)
    
    Z = np.zeros_like(X)
    
    # Map for coded variable names
    var_map = {'X1': 0, 'X2': 0, 'X3': 0, 'X4': 0}

    # Calculate predicted yield (Z) for the grid
    for i in range(X.shape[0]):
        for j in range(X.shape[1]):
            
            # Reset and map the grid values to the correct coded factor
            temp_map = var_map.copy()
            temp_map[X_var_code] = X[i, j]
            temp_map[Y_var_code] = Y[i, j]
            
            # The other two factors remain at 0 (fixed)
            
            Z[i, j] = predict_yield_coded(temp_map['X1'], temp_map['X2'], temp_map['X3'], temp_map['X4'], coeffs)

    x_label_actual = get_actual_label(X_var_code)
    y_label_actual = get_actual_label(Y_var_code)
    fixed_label = ', '.join([get_actual_label(code) for code in fixed_var_codes])
    title_suffix = f" (Fixed at Center: {fixed_label})"
    
    # 2D Contour Plot
    fig_2d = plt.figure(figsize=(8, 6))
    contour = plt.contourf(X, Y, Z, cmap='viridis', levels=20)
    plt.colorbar(contour, label='Predicted Oil Yield (%)')
    plt.xlabel(f'{x_label_actual} ({X_var_code} Coded)')
    plt.ylabel(f'{y_label_actual} ({Y_var_code} Coded)')
    plt.title(f'2D Contour Plot: {X_var_code} vs {Y_var_code}' + title_suffix)
    plot_2d_filename = f'contour_{X_var_code}_{Y_var_code}.png'
    fig_2d.savefig(plot_2d_filename, dpi=300)
    plt.close(fig_2d)
    
    # 3D Surface Plot
    fig_3d = plt.figure(figsize=(10, 8))
    ax = fig_3d.add_subplot(111, projection='3d')
    ax.plot_surface(X, Y, Z, cmap='viridis')
    ax.set_xlabel(f'{x_label_actual} ({X_var_code} Coded)')
    ax.set_ylabel(f'{y_label_actual} ({Y_var_code} Coded)')
    ax.set_zlabel('Predicted Oil Yield (%)')
    ax.set_title(f'3D Surface Plot: {X_var_code} vs {Y_var_code}' + title_suffix)
    plot_3d_filename = f'surface_{X_var_code}_{Y_var_code}.png'
    fig_3d.savefig(plot_3d_filename, dpi=300)
    plt.close(fig_3d)
    
    return plot_2d_filename, plot_3d_filename

# Factor combinations to plot (fixing the other two at center point 0)
plot_combinations = [
    ('X1', 'X2', ['X3', 'X4']),  # MC vs HTemp
    ('X3', 'X4', ['X1', 'X2']),  # HTime vs SET
    ('X2', 'X4', ['X1', 'X3']),  # HTemp vs SET
]

plot_files_list = []
for X_var, Y_var, fixed_vars in plot_combinations:
    files = generate_rsm_plots(X_var, Y_var, fixed_vars, coefficients)
    plot_files_list.extend(files)

# --- 5. Document Generation ---

def generate_word_document(results, opt_results, plot_files):
    
    document = Document()
    document.add_heading('Response Surface Methodology Analysis: Jatropha Oil Extraction', 0)
    
    # --- Optimization Results ---
    document.add_heading('Optimal Factor Setting and Predicted Yield', level=1)
    
    document.add_paragraph(
        f"The numerical optimization using the fitted regression model found the maximum predicted oil yield of "
        f"{opt_results['Y_opt']:.4f} %."
    )
    
    document.add_paragraph("The corresponding optimal factor settings (validating the experimental results) are:")
    
    document.add_paragraph(
        f"  • Moisture Content (MC): {opt_results['MC_opt']:.2f} % wet basis"
    )
    document.add_paragraph(
        f"  • Heating Temperature (HT): {opt_results['HTemp_opt']:.2f} °C"
    )
    document.add_paragraph(
        f"  • Heating Time (Ht): {opt_results['HTime_opt']:.2f} mins"
    )
    document.add_paragraph(
        f"  • Soxhlet Extraction Time (SET): {opt_results['SET_opt']:.2f} mins"
    )

    # --- Model Summary and ANOVA ---
    document.add_heading('Model Summary and ANOVA', level=1)
    
    # Add ANOVA and Summary as formatted text block
    # Note: statsmodels summary contains R-squared, coefficients, and ANOVA table (as a block)
    summary_text = results.summary().as_text()
    
    document.add_paragraph("Regression Coefficients, R-squared, and P-values:")
    document.add_paragraph(summary_text)

    # --- Plots ---
    document.add_heading('2D Contour and 3D Surface Plots', level=1)
    
    for file in plot_files:
        document.add_heading(file.replace('.png', '').replace('_', ' ').title(), level=2)
        document.add_picture(file, width=Inches(5.5))
        
    document.save('RSM_Analysis_Report.docx')
    print("\nSuccessfully created RSM_Analysis_Report.docx")
    
    # --- Cleanup Temporary Files ---
    for file in plot_files:
        try:
            os.remove(file)
        except Exception as e:
            print(f"Warning: Could not delete temporary file {file}. Error: {e}")
            
    print("Temporary plot images deleted.")

# Execute the document generation
generate_word_document(results, opt_results_dict, plot_files_list)

