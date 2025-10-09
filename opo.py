from docx import Document
from docx.shared import Pt

# Create a Word document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
doc.add_heading("Roller Mill Drive System Analysis", 0)

# Assumptions
doc.add_heading("1. Assumptions", level=1)
assumptions = [
    "Driver (hammer) speed N1 = 700 rpm.",
    "Roller speed N2 = 280 rpm (given).",
    "Hammer pulley diameter (driver) D1 = 10 in = 254.0 mm.",
    "Driven pulley diameter (roller) D2 = 25 in = 635.0 mm (calculated from speeds).",
    "Roller geometry: 55 mm diameter, 130 mm length, 1 mm gap.",
    "Bulk density of grain/flour: 750 kg/m³.",
    "Fill (nip) factor: f = 0.6.",
    "Assumed pulley center distance: C = 500 mm.",
    "Specific energy for roller milling: 20 kWh/ton.",
    "Steel density: 7850 kg/m³."
]
for item in assumptions:
    doc.add_paragraph(item, style='List Bullet')

# Pulley diameters
doc.add_heading("2. Pulley Diameters (mm)", level=1)
doc.add_paragraph("Driver (hammer) pulley: 254.0 mm")
doc.add_paragraph("Driven (roller) pulley: 635.0 mm")

# Belt length
doc.add_heading("3. Belt Length", level=1)
doc.add_paragraph("Calculated using open-belt formula with center distance C = 500 mm:")
doc.add_paragraph("Belt length L ≈ 2469 mm")

# Wrap angles
doc.add_heading("4. Belt Wrap Angles", level=1)
doc.add_paragraph("Small pulley (driver): 135.21°")
doc.add_paragraph("Large pulley (driven): 224.79°")

# Peripheral speed
doc.add_heading("5. Roller Peripheral Speed", level=1)
doc.add_paragraph("Roller surface speed ≈ 1.61 m/s")

# Theoretical capacity
doc.add_heading("6. Theoretical Roller Mill Capacity", level=1)
doc.add_paragraph("Ideal mass flow rate (nip-fill): ≈ 169.8 kg/h")

# Power and torque
doc.add_heading("7. Power and Torque", level=1)
doc.add_paragraph("Estimated shaft power: ≈ 3.396 kW")
doc.add_paragraph("Torque at 280 rpm: ≈ 115.84 N·m")

# Roller mass and centrifugal force
doc.add_heading("8. Roller Mass and Centrifugal Force", level=1)
doc.add_paragraph("Roller mass (solid steel): ≈ 2.425 kg")
doc.add_paragraph("Centrifugal force at 280 rpm: ≈ 57.32 N")

# Summary
doc.add_heading("9. Summary", level=1)
summary_points = [
    "Driver pulley D1 = 254.0 mm",
    "Driven pulley D2 = 635.0 mm",
    "Belt length ≈ 2469 mm",
    "Wrap angle (small pulley) ≈ 135.21°",
    "Roller speed ≈ 1.61 m/s",
    "Theoretical capacity ≈ 169.8 kg/h",
    "Power requirement ≈ 3.396 kW",
    "Torque ≈ 115.84 N·m",
    "Roller mass ≈ 2.425 kg",
    "Centrifugal force ≈ 57.32 N"
]
for item in summary_points:
    doc.add_paragraph(item, style='List Bullet')

# Notes
doc.add_heading("10. Notes and Recommendations", level=1)
notes = [
    "If 'A40' was intended to represent a belt length (40 in), it's too short for these pulley sizes.",
    "Wrap angle of 135° is generally acceptable; if it drops below 120°, consider a tensioner/idler.",
    "Practical throughput is typically 10–60% of theoretical capacity.",
    "Use a motor with at least 1.25× service factor. A 4–5 kW motor is recommended.",
    "To improve accuracy, measure actual motor power and update calculations."
]
for item in notes:
    doc.add_paragraph(item, style='List Number')

# Save the document
doc.save("roller_mill_drive_analysis.docx")

