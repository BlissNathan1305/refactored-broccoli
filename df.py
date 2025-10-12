# test_cover_docx.py
# Minimal, robust test: create a cover image and embed it into a .docx
# Works with Pillow >=10 and falls back if fonts are missing.

from PIL import Image, ImageDraw, ImageFont, ImageFilter
from docx import Document
from docx.shared import Inches
import os

COVER = "cover_test.png"
OUTDOC = "test_book.docx"

def create_cover(path=COVER, w=1600, h=2300):
    # Simple gradient + centered text using textbbox (Pillow 10+)
    top = (20, 50, 120)
    bottom = (40, 140, 110)
    img = Image.new("RGB", (w,h), top)
    draw = ImageDraw.Draw(img)

    for y in range(h):
        t = y / (h-1)
        r = int(top[0]*(1-t) + bottom[0]*t)
        g = int(top[1]*(1-t) + bottom[1]*t)
        b = int(top[2]*(1-t) + bottom[2]*t)
        draw.line([(0,y),(w,y)], fill=(r,g,b))

    # geometric accent
    for i in range(0, w, 150):
        draw.line([(i, 0), (i - 300, h)], fill=(255,255,255,20), width=2)

    title = "DATA ANALYSIS\nUSING PYTHON & R"
    author = "By Samuel Blessed Nathaniel"

    # load font safely
    def try_font(name, size):
        try:
            return ImageFont.truetype(name, size)
        except:
            return ImageFont.load_default()

    title_font = try_font("DejaVuSans-Bold.ttf", 64)
    author_font = try_font("DejaVuSans.ttf", 28)

    # draw centered multiline title
    lines = title.split("\n")
    total_h = 0
    line_sizes = []
    for line in lines:
        bbox = draw.textbbox((0,0), line, font=title_font)
        w_line = bbox[2]-bbox[0]; h_line = bbox[3]-bbox[1]
        line_sizes.append((w_line, h_line))
        total_h += h_line + 8
    start_y = int(h*0.32 - total_h//2)

    cur_y = start_y
    for i, line in enumerate(lines):
        w_line, h_line = line_sizes[i]
        x = (w - w_line)//2
        draw.text((x, cur_y), line, font=title_font, fill=(255,255,255))
        cur_y += h_line + 8

    # author
    bbox = draw.textbbox((0,0), author, font=author_font)
    w_a = bbox[2]-bbox[0]; h_a = bbox[3]-bbox[1]
    draw.text(((w-w_a)//2, cur_y + 30), author, font=author_font, fill=(245,245,245))

    img = img.filter(ImageFilter.SMOOTH)
    img.save(path)
    return path

def build_docx(cover_path, outpath=OUTDOC):
    doc = Document()
    # Insert cover full width (approx)
    try:
        doc.add_picture(cover_path, width=Inches(6))
    except Exception as e:
        # fallback: try smaller width
        doc.add_picture(cover_path, width=Inches(4))
    doc.add_page_break()
    doc.add_heading("Test Book", 0)
    doc.add_paragraph("This is a minimal generated document. The cover above is the generated image.")
    doc.save(outpath)
    return outpath

if __name__ == "__main__":
    print("Creating cover image...")
    cover_path = create_cover()
    print("Cover saved to:", cover_path)
    print("Creating docx and embedding cover...")
    out = build_docx(cover_path)
    print("Done. Document saved to:", out)
    print("If you can't see the cover in previews, open the file in LibreOffice or MS Word on a desktop.")
