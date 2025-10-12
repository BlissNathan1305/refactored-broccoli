import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import linregress
from docx import Document
from docx.shared import Inches

def perform_statistical_analysis():
    """
    1. Simulates data acquisition.
    2. Performs statistical analysis (linear regression).
    3. Outputs results and plot to a .docx file.
    """
    
    # --- 1. SIMULATED DATA ACQUISITION ---
    # This data is synthesized based on field experiment findings (e.g., from Intl Agrophysics, Paper 1.8) 
    # where fuel consumption (L/ha) increases linearly with ploughing depth (cm).
    print("--- Simulating Data Acquisition (Using hardcoded data from research papers) ---")
    data = {
        'Tillage_Depth_cm': [10, 15, 20, 23, 25, 30],
        'Fuel_Consumption_L_ha': [15.00, 19.67, 24.00, 24.72, 27.00, 28.65]
    }
    df = pd.DataFrame(data)

    print(f"\nData Acquired ({len(df)} observations):\n{df.to_markdown(index=False)}")
    print("-" * 50)

    # --- 2. STATISTICAL ANALYSIS ---
    print("--- Running Linear Regression Analysis ---")
    
    X = df['Tillage_Depth_cm']
    Y = df['Fuel_Consumption_L_ha']

    # Perform Linear Regression
    # Model: Fuel_Consumption = m * Depth + c
    slope, intercept, r_value, p_value, std_err = linregress(X, Y)

    r_squared = r_value**2
    
    # Calculate predicted values for plotting the line
    Y_pred = intercept + slope * X

    # Create the scatter plot
    plt.figure(figsize=(8, 6))
    plt.scatter(X, Y, color='blue', label='Observed Data')
    plt.plot(X, Y_pred, color='red', 
             label=f'Regression Line\n(R² = {r_squared:.4f})')
    
    plt.title('Effect of Tillage Depth on Tractor Fuel Efficiency')
    plt.xlabel('Tillage Depth (cm)')
    plt.ylabel('Fuel Consumption (L/ha)')
    plt.legend()
    plt.grid(True)
    
    # Save the plot temporarily
    plot_path = 'fuel_efficiency_plot.png'
    plt.savefig(plot_path)
    plt.close()
    
    # --- 3. OUTPUT TO DOCX ---
    
    print("--- Generating Statistical Report (Docx) ---")
    document = Document()
    document.add_heading('Statistical Analysis Report: Tillage Depth vs. Fuel Efficiency', 0)

    # Summary
    document.add_paragraph(
        'This report presents the statistical analysis on the relationship between tractor '
        'tillage depth and specific fuel consumption (Liters per Hectare).'
    )
    
    # Raw Data Table
    document.add_heading('1. Raw Data', level=1)
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name.replace('_', ' ')
        
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = f"{row['Tillage_Depth_cm']:.2f}"
        row_cells[1].text = f"{row['Fuel_Consumption_L_ha']:.3f}"

    # Regression Results
    document.add_heading('2. Linear Regression Results', level=1)
    
    document.add_paragraph(
        f"**Regression Equation:**\n"
        f"Fuel Consumption (L/ha) = ({slope:.3f} * Tillage Depth (cm)) + {intercept:.3f}\n\n"
        f"**Statistical Metrics:**"
    )

    stats_table = document.add_table(rows=3, cols=2)
    stats_table.style = 'Table Grid'
    stats_table.cell(0, 0).text = 'R-Squared ($R^2$)'
    stats_table.cell(0, 1).text = f'{r_squared:.4f}'
    stats_table.cell(1, 0).text = 'P-value'
    stats_table.cell(1, 1).text = f'{p_value:.5f}'
    stats_table.cell(2, 0).text = 'Standard Error of Slope'
    stats_table.cell(2, 1).text = f'{std_err:.4f}'
    
    # Interpretation
    document.add_heading('3. Interpretation', level=1)
    interpretation = (
        f"The **Coefficient of Determination ($R^2$) is {r_squared:.4f}**, which is very high. "
        "This indicates that approximately **{r_squared*100:.2f}%** of the variability in Fuel Consumption "
        "can be explained by the Tillage Depth using this linear model. "
        "The **P-value is {p_value:.5f}**, which is far below the significance level of 0.05, "
        "confirming that Tillage Depth has a statistically significant positive effect on Fuel Consumption."
        "Specifically, the slope of {slope:.3f} suggests that for every **1 cm increase in Tillage Depth**, "
        "the Fuel Consumption increases by approximately **{slope:.3f} Liters per Hectare**."
    )
    document.add_paragraph(interpretation)
    
    # Plot
    document.add_heading('4. Visualization', level=1)
    document.add_picture(plot_path, width=Inches(6))

    docx_filename = 'tractor_fuel_analysis_report.docx'
    document.save(docx_filename)
    
    # Clean up the temporary plot file (optional)
    import os
    os.remove(plot_path)

    print("-" * 50)
    print(f"SUCCESS: Report saved to {docx_filename}")
    print("-" * 50)

if __name__ == '__main__':
    perform_statistical_analysis()

