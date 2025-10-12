# ===============================================
# PRACTICAL DATA ANALYSIS WITH PYTHON AND R
# A Hands-On Guide (Full Book)
# Author: Samuel Blessed Nathaniel
# Edition: 2025
# ===============================================

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import textwrap

# ------------------------------------------------
# Create Document
# ------------------------------------------------
doc = Document()

def add_title_page():
    doc.add_section(start_type=0)
    p = doc.add_paragraph("\n\n\n\n")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRACTICAL DATA ANALYSIS\nWITH PYTHON AND R\n")
    run.font.size = Pt(36)
    run.bold = True
    run.font.name = 'Times New Roman'

    run = p.add_run("A HANDS-ON GUIDE\n\n")
    run.font.size = Pt(20)

    run = p.add_run("By Samuel Blessed Nathaniel\n")
    run.font.size = Pt(18)

    run = p.add_run(f"\n{date.today().year} Edition\n")
    run.font.size = Pt(14)
    doc.add_page_break()

def add_table_of_contents():
    doc.add_heading("Table of Contents", level=1)
    chapters = [
        "Part I – Introduction to Data Analysis",
        "Part II – Working with Data",
        "Part III – Analysis and Modeling",
        "Part IV – Reporting, Automation & Exporting Results",
        "Chapter 1: Introduction to Data Analysis",
        "Chapter 2: Understanding Data",
        "Chapter 3: Data Collection and Importing",
        "Chapter 4: Data Cleaning and Preprocessing",
        "Chapter 5: Data Exploration (EDA)",
        "Chapter 6: Statistical Analysis in Python and R",
        "Chapter 7: Data Visualization Masterclass",
        "Chapter 8: Predictive Modeling and Machine Learning",
        "Chapter 9: Communicating Results Effectively",
        "Chapter 10: Creating Reports in Python",
        "Chapter 11: Creating Reports in R",
        "Chapter 12: Building Dashboards",
        "Chapter 13: Data Analysis Project Workflow",
        "Chapter 14: Best Practices & Ethics"
    ]
    for ch in chapters:
        para = doc.add_paragraph(ch)
        para.style = 'List Number'
    doc.add_page_break()

def add_chapter(title, content):
    doc.add_heading(title, level=1)
    for paragraph in content.strip().split("\n\n"):
        doc.add_paragraph(paragraph.strip())
    doc.add_page_break()

def add_code_block(code):
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)

# ------------------------------------------------
# Build Book
# ------------------------------------------------
add_title_page()
add_table_of_contents()

# --- Chapter 1 ---
add_chapter("Chapter 1: Introduction to Data Analysis", """
Data analysis is the process of collecting, transforming, and interpreting data to extract meaningful insights.
It is the foundation of data-driven decision-making in every modern field — from healthcare to finance, business,
and public policy.

In this book, you’ll learn how to analyze data practically using Python and R — two of the most powerful tools
for analytics, visualization, and reporting.
""")

# --- Chapter 2 ---
add_chapter("Chapter 2: Understanding Data", """
Data is information in raw form — numbers, text, images, or observations collected from various sources.
Before analysis, understanding data structure, type, and quality is critical.

There are two main categories:
- **Quantitative data**: numerical, measurable (e.g., sales, temperature).
- **Qualitative data**: categorical, descriptive (e.g., gender, color, brand).

Python and R can handle all these types using data structures like lists, dictionaries, and data frames.
""")

# --- Chapter 3 ---
add_chapter("Chapter 3: Data Collection and Importing", """
Data collection involves gathering raw information from reliable sources — databases, files, surveys, or APIs.

Example: Importing data in Python and R.
""")

add_code_block(textwrap.dedent("""
# Python
import pandas as pd
data = pd.read_csv("sales.csv")
"""))

add_code_block(textwrap.dedent("""
# R
library(readr)
data <- read_csv("sales.csv")
"""))

# --- Chapter 4 ---
add_chapter("Chapter 4: Data Cleaning and Preprocessing", """
Data cleaning ensures accuracy by fixing missing values, duplicates, or incorrect formats.

Common steps:
1. Handle missing data
2. Remove duplicates
3. Convert data types
4. Handle outliers
5. Standardize formats
""")

add_code_block(textwrap.dedent("""
# Python example
df.dropna(inplace=True)
df = df.drop_duplicates()
"""))

# --- Chapter 5 ---
add_chapter("Chapter 5: Data Exploration (EDA)", """
Exploratory Data Analysis (EDA) helps you understand the patterns, distributions, and relationships in your data.

Typical steps:
- Compute descriptive statistics
- Visualize distributions (histograms, boxplots)
- Identify correlations
""")

add_code_block(textwrap.dedent("""
# Python
df.describe()
df.corr()
"""))

add_code_block(textwrap.dedent("""
# R
summary(df)
cor(df)
"""))

# --- Chapter 6 ---
add_chapter("Chapter 6: Statistical Analysis in Python and R", """
Statistical analysis involves applying mathematical formulas to interpret and validate trends.

Examples:
- Hypothesis testing
- Correlation analysis
- ANOVA
- Regression models
""")

# --- Chapter 7 ---
add_chapter("Chapter 7: Data Visualization Masterclass", """
Data visualization transforms insights into compelling stories through charts and dashboards.
Python uses Matplotlib, Seaborn, and Plotly.
R uses ggplot2 and lattice.
""")

add_code_block(textwrap.dedent("""
# Python
import matplotlib.pyplot as plt
plt.bar(df['Region'], df['Sales'])
plt.show()
"""))

# --- Chapter 8 ---
add_chapter("Chapter 8: Predictive Modeling and Machine Learning", """
Machine learning allows data to make predictions. Typical models:
- Linear Regression
- Decision Trees
- Random Forests
- K-Means Clustering
""")

# --- Chapter 9 ---
add_chapter("Chapter 9: Communicating Results Effectively", """
Data storytelling makes your insights actionable. Focus on:
- Simplicity
- Accuracy
- Visualization clarity
- Contextual recommendations
""")

# --- Chapter 10 ---
add_chapter("Chapter 10: Creating Reports in Python", """
Python can export professional reports in DOCX, PDF, and HTML formats using libraries like:
- python-docx
- reportlab
- Jupyter Notebook
- Streamlit

Example:
""")

add_code_block(textwrap.dedent("""
from docx import Document
doc = Document()
doc.add_heading("Sales Analysis Report", 0)
doc.add_paragraph("Generated using Python.")
doc.save("Report.docx")
"""))

# --- Chapter 11 ---
add_chapter("Chapter 11: Creating Reports in R", """
R Markdown and Officer packages help create automated reports in DOCX or PDF format.

Example:
""")

add_code_block(textwrap.dedent("""
library(officer)
doc <- read_docx()
doc <- body_add_par(doc, "Sales Report", style="heading 1")
print(doc, target="Sales_Report.docx")
"""))

# --- Chapter 12 ---
add_chapter("Chapter 12: Building Dashboards", """
Dashboards combine visuals, summaries, and interactivity.
Python: Streamlit, Dash
R: Shiny
""")

# --- Chapter 13 ---
add_chapter("Chapter 13: Data Analysis Project Workflow", """
A professional workflow includes:
1. Define problem
2. Collect data
3. Clean and preprocess
4. Explore and visualize
5. Model and evaluate
6. Report and automate
""")

# --- Chapter 14 ---
add_chapter("Chapter 14: Best Practices & Ethics", """
Be ethical and transparent. Protect privacy, avoid bias, and document assumptions.
A good analyst ensures reproducibility and honesty in data interpretation.
""")

# ------------------------------------------------
# Save Book
# ------------------------------------------------
doc.save("Practical_Data_Analysis_with_Python_and_R.docx")
print("✅ Book successfully generated as 'Practical_Data_Analysis_with_Python_and_R.docx'")
