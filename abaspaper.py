# save as generate_dryer_report.py
import os
import math
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# -------------------------
# 1) Data (extracted from uploaded file: Appendix B)
# -------------------------
# Times in minutes (upper): 0,60,...,780 (14 points)
times_upper = np.array([0, 60, 120, 180, 240, 300, 360, 420, 480, 540, 600, 660, 720, 780])
# Average moisture content (% wb) for upper tray (from Appendix B.1)
mc_upper = np.array([67, 64, 63, 60, 56, 45, 37, 36, 21, 16, 11, 7, 3, 0], dtype=float)

# Times in minutes (lower): 0,60,120,...,1020 (19 points)
times_lower = np.array([0, 60, 120, 180, 240, 300, 360, 420, 480, 540, 600, 660, 720, 780, 840, 900, 960, 1020, 1140])
# Average moisture content (% wb) for lower tray (from Appendix B.1, truncated/aligned to match table)
mc_lower = np.array([68, 69, 65, 60, 55, 51, 43, 40, 38, 30, 25, 22, 18, 14, 10, 7, 3, 0, 0], dtype=float)

# Times and MC for open-air drying (0..1200 at 60-min intervals; using available Appendix B values)
times_open = np.array([0,60,120,180,240,300,360,480,540,600,660,720,780,840,900,960,1020,1080,1140,1200])
mc_open = np.array([52,50,49,47,46,45,44,42,31,26,24,22,19,16,13,9,7,3,3,0], dtype=float)

# -------------------------
# 2) Compute drying rates (g H2O / min) from MC change per minute
# For simplicity, assume mass basis of sample is proportional to MC differences (calculate derivative of MC).
# -------------------------
def compute_drying_rate(time_min, mc_percent):
    # Convert to water mass units if initial sample mass known; here we report rate in %wb per min.
    dt = np.diff(time_min)
    dmc = np.diff(mc_percent)
    rate = dmc / dt  # %wb per minute (negative for decreasing moisture)
    # return rate center-aligned (time at midpoint of interval)
    t_mid = time_min[:-1] + dt/2
    return t_mid, rate

t_mid_u, rate_u = compute_drying_rate(times_upper, mc_upper)
t_mid_l, rate_l = compute_drying_rate(times_lower, mc_lower)
t_mid_o, rate_o = compute_drying_rate(times_open, mc_open)

# Convert rate to g_H2O/min equivalent if a sample initial mass is known.
# The source file reported drying rates in g H2O/min calculated from mass changes.
# If required, replace with actual mass -> water mass conversions.

# -------------------------
# 3) Basic statistics
# -------------------------
def summary_stats(arr):
    return {
        'mean': float(np.nanmean(arr)),
        'std': float(np.nanstd(arr, ddof=1)),
        'min': float(np.nanmin(arr)),
        'max': float(np.nanmax(arr)),
    }

stats = {
    'upper_mc': summary_stats(mc_upper),
    'lower_mc': summary_stats(mc_lower),
    'open_mc': summary_stats(mc_open),
    'upper_temp_mean': 40.24,  # from file (°C)
    'ambient_temp_mean': 32.77,  # from file (°C)
}

# -------------------------
# 4) Plots
# -------------------------
os.makedirs('figures', exist_ok=True)

# Moisture vs time
plt.figure(figsize=(7,4))
plt.plot(times_upper/60, mc_upper, marker='o', label='Upper tray')
plt.plot(times_lower/60, mc_lower, marker='s', label='Lower tray')
plt.plot(times_open/60, mc_open, marker='^', label='Open air')
plt.xlabel('Time (hours)')
plt.ylabel('Moisture content (% wb)')
plt.title('Moisture content vs drying time')
plt.grid(True)
plt.legend()
plt.tight_layout()
plt.savefig('figures/moisture_vs_time.png', dpi=200)
plt.close()

# Drying rate vs time (%wb/min)
plt.figure(figsize=(7,4))
plt.plot(t_mid_u/60, -rate_u, marker='o', label='Upper tray (rate)')
plt.plot(t_mid_l/60, -rate_l, marker='s', label='Lower tray (rate)')
plt.plot(t_mid_o/60, -rate_o, marker='^', label='Open air (rate)')
plt.xlabel('Time (hours)')
plt.ylabel('Drying rate (%wb per min) (absolute)')
plt.title('Drying rate vs time (absolute values)')
plt.grid(True)
plt.legend()
plt.tight_layout()
plt.savefig('figures/drying_rate_vs_time.png', dpi=200)
plt.close()

# -------------------------
# 5) Build docx report
# -------------------------
doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

# Title
h = doc.add_heading('Development and Evaluation of a Direct Passive Polycarbonate Cylindrical Solar Dryer', level=1)
h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Authors (placeholder)
p = doc.add_paragraph('Utit, I. I.; Supervisor: Dr. David Onwe\n\n', style='Normal')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Abstract
doc.add_heading('Abstract', level=2)
doc.add_paragraph(
    "This report reproduces and analyses experimental drying data (crayfish) from a cylindrical polycarbonate solar dryer. "
    "The full dataset and methods were taken from the supplied project document (design, raw data and calculations). "
    "The document contains moisture curves, drying rates, summary statistics and recommendations."
)

# Methods summary
doc.add_heading('Materials and Methods', level=2)
doc.add_paragraph(
    "Design: cylindrical polycarbonate dryer (volume 52,297 cm3), two tray levels. "
    "Measurements: mass (three replicates), ambient and dryer temperature, relative humidity. "
    "Data source: project file (Appendix A/B)."
)

# Results: add table of moisture vs time for upper tray
doc.add_heading('Results', level=2)
doc.add_paragraph('Moisture content – Upper tray (mean of 3 replicates):')
df_upper = pd.DataFrame({'Time_min': times_upper, 'MC_pct_wb': mc_upper})
doc.add_paragraph(df_upper.to_string(index=False))
# Insert figure
doc.add_picture('figures/moisture_vs_time.png', width=Inches(6.0))
doc.add_paragraph('Figure 1. Moisture content vs drying time (upper, lower & open air).')

# Add drying rate figure
doc.add_picture('figures/drying_rate_vs_time.png', width=Inches(6.0))
doc.add_paragraph('Figure 2. Drying rate vs time (absolute).')

# Summary statistics
doc.add_heading('Summary statistics', level=3)
for k,v in stats.items():
    if isinstance(v, dict):
        doc.add_paragraph(f"{k}: mean={v['mean']:.2f}, std={v['std']:.2f}, min={v['min']:.2f}, max={v['max']:.2f}")
    else:
        doc.add_paragraph(f"{k}: {v}")

# Discussion & Conclusion
doc.add_heading('Discussion and Conclusion', level=2)
doc.add_paragraph(
    "The dryer reached internal temperatures above ambient, reduced drying time relative to open air, "
    "and yielded typical drying curves with initial faster rates followed by falling-rate periods. "
    "Recommendations include increasing absorber area and testing venting/ chimney options."
)

# Save docx
out_path = 'solar_dryer_report.docx'
doc.save(out_path)
print(f"Report and figures created. Open '{out_path}' and the 'figures' folder.")

# End of script
