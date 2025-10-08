import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# Create output folder
output_folder = "soil_charts"
os.makedirs(output_folder, exist_ok=True)

# Soil data
data = {
    'Location': ['Nung Uyo Idoro', 'Nung Uyo Idoro', 'Ekom Iman', 'Ekom Iman',
                 'Afaha Idoro', 'Afaha Idoro', 'Ikot Idaha', 'Ikot Idaha'],
    'Depth': ['0-50', '50-100'] * 4,
    'pH': [7.08, 7.26, 7.10, 7.37, 7.30, 7.62, 7.82, 5.63],
    'OC (%)': [2.19, 1.80, 1.02, 0.96, 2.63, 1.54, 2.23, 0.88],
    'T/N (%)': [0.09, 0.08, 0.04, 0.04, 0.11, 0.07, 0.10, 0.22],
    'Av.P (mg/kg)': [62.0, 54.0, 56.666, 60.666, 53.999, 43.332, 54.0, 86.666],
    'Ca': [4.56, 4.08, 4.56, 4.80, 4.32, 6.24, 5.76, 3.34],
    'Mg': [2.88, 2.06, 2.72, 2.96, 2.68, 4.16, 3.44, 1.94],
    'K': [0.06, 0.06, 0.03, 0.07, 0.03, 0.08, 0.03, 0.02],
    'Na': [0.05, 0.03, 0.09, 0.05, 0.09, 0.13, 0.06, 0.07],
    'BSAT (%)': [81.79, 83.84, 73.70, 67.24, 72.95, 71.06, 79.47, 77.74]
}

df = pd.DataFrame(data)

# Function to plot and save chart
def plot_and_save(property_name):
    fig, ax = plt.subplots(figsize=(10, 6))
    for depth in ['0-50', '50-100']:
        subset = df[df['Depth'] == depth]
        ax.bar(subset['Location'], subset[property_name], label=f'Depth {depth}', alpha=0.7)
    
    ax.set_title(f'{property_name} Across Locations and Depths')
    ax.set_ylabel(property_name)
    ax.set_xlabel('Location')
    ax.legend()
    plt.xticks(rotation=45)
    plt.tight_layout()
    
    filename = f"{output_folder}/{property_name.replace(' ', '_')}.png"
    plt.savefig(filename)
    plt.close()
    return filename

# Create Word document
doc = Document()
doc.add_heading('Soil Physicochemical Properties Across Locations', 0)

# Add charts and captions
properties = ['pH', 'OC (%)', 'T/N (%)', 'Av.P (mg/kg)', 'Ca', 'Mg', 'K', 'Na', 'BSAT (%)']
for prop in properties:
    chart_path = plot_and_save(prop)
    doc.add_heading(prop, level=1)
    doc.add_picture(chart_path, width=Inches(6))
    doc.add_paragraph(f"Bar chart showing {prop} variation across locations at depths 0–50 cm and 50–100 cm.")

# Save document
doc.save("Soil_Properties_Report.docx")
print("Document exported as 'Soil_Properties_Report.docx'")
