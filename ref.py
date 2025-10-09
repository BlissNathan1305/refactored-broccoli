from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()

# Set default font style
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
doc.add_heading('Academic Discussion: Jatropha Oil Extraction Analysis', level=1)

# Section 1: Regression Model Summary
doc.add_heading('1. Regression Model Summary', level=2)
doc.add_paragraph(
    "The regression model developed to predict oil yield from Jatropha seeds shows an R-squared value of 0.5906, "
    "indicating that approximately 59.06% of the variability in oil yield is explained by the model. However, the "
    "Adjusted R-squared is significantly lower at 0.1519, suggesting that the inclusion of multiple predictors may not "
    "be contributing meaningfully to the model’s explanatory power. This discrepancy often points to overfitting, where "
    "the model captures noise rather than signal.\n\n"
    "RMSE (2.0827) and MSE (4.3377) reflect the average prediction error magnitude. For engineering applications, especially "
    "in process optimization, these values should be minimized to ensure reliable predictions. The relatively high RMSE implies "
    "that the model may not be sufficiently precise for deployment without further refinement.\n\n"
    "Implication: The model may benefit from variable selection techniques or regularization (e.g., Lasso or Ridge regression) "
    "to improve generalizability and reduce overfitting."
)

# Section 2: Model Coefficients
doc.add_heading('2. Model Coefficients', level=2)
doc.add_paragraph(
    "The regression equation includes linear, quadratic, and interaction terms for four independent variables (X1 to X4), "
    "likely representing:\n- X1: Moisture Content (%)\n- X2: Heating Temperature (°C)\n- X3: Heating Time (min)\n- X4: Extraction Time (min)\n\n"
    "Key Observations:\n"
    "- X1 (Moisture Content) has a strong negative linear coefficient (-1.0142) and negative quadratic and interaction terms, indicating that higher moisture content consistently reduces oil yield.\n"
    "- X2 (Heating Temperature) and X4 (Extraction Time) show positive linear effects, suggesting that higher temperatures and longer extraction durations enhance yield.\n"
    "- X3 (Heating Time) has a moderate positive linear effect but a negative quadratic term, implying diminishing returns or possible degradation at extended heating durations.\n"
    "- Interaction terms such as X1 X3 (-1.4000) and X1 X4 (-1.6900) are notably negative, revealing that moisture content negatively interacts with both heating and extraction time.\n\n"
    "Engineering Insight: These coefficients provide a nuanced understanding of process dynamics. Engineers can use this to fine-tune parameters, avoiding combinations that lead to reduced efficiency."
)

# Section 3: ANOVA
doc.add_heading('3. Analysis of Variance (ANOVA)', level=2)
doc.add_paragraph(
    "The ANOVA table reveals:\n"
    "- F-value of 1.5454 and P-value of 0.2064 for the regression model indicate that the model is not statistically significant at the conventional 0.05 level.\n"
    "- Residual Mean Square (8.6755) is relatively high, reinforcing the earlier concern about prediction error.\n\n"
    "Implication for Engineering Research: While the model structure is comprehensive, its statistical insignificance calls for reevaluation. Possible remedies include increasing sample size, refining variable definitions, or exploring nonlinear modeling techniques such as artificial neural networks or support vector regression."
)

# Section 4: Response Surface Plots
doc.add_heading('4. Response Surface Plots', level=2)
doc.add_paragraph(
    "Though the plots are not shown, the report lists four key bivariate relationships:\n"
    "1. Moisture Content vs Heating Temperature\n"
    "2. Moisture Content vs Extraction Time\n"
    "3. Heating Temperature vs Extraction Time\n"
    "4. Heating Time vs Extraction Time\n\n"
    "These plots are essential for visualizing interaction effects and curvature in the response surface. They help identify saddle points, ridges, and valleys in the yield landscape, guiding engineers toward optimal operating zones.\n\n"
    "Recommendation: Future work should include 3D surface plots and contour maps to facilitate intuitive understanding and support decision-making in process design."
)

# Section 5: Optimization Results
doc.add_heading('5. Optimization Results', level=2)
doc.add_paragraph(
    "The optimization algorithm predicts a maximum oil yield of 37.92% under the following conditions:\n"
    "- Moisture Content: 6.00%\n"
    "- Heating Temperature: 90.00°C\n"
    "- Heating Time: 35.00 min\n"
    "- Extraction Time: 180.00 min\n\n"
    "This is a significant improvement over the experimental maximum of 26.70%, achieved under different conditions (MC=12%, HT=90°C, Ht=30min, SET=180min).\n\n"
    "Discussion:\n"
    "- The 11.22% increase in predicted yield underscores the value of optimization techniques in engineering design.\n"
    "- The lower moisture content in the optimized setup aligns with the regression findings.\n"
    "- The extended heating time may allow better thermal penetration and oil release, though care must be taken to avoid thermal degradation.\n\n"
    "Engineering Application: These results can inform pilot-scale trials and industrial scale-up, with the caveat that model validation is essential before implementation."
)

# Section 6: Regression Model Equation
doc.add_heading('6. Regression Model Equation', level=2)
doc.add_paragraph(
    "The full regression equation encapsulates the complex interplay of linear, quadratic, and interaction effects. For engineering use, this equation can be embedded into simulation tools or control systems to predict yield under varying conditions.\n\n"
    "However, given the model’s statistical limitations, it should be treated as a first-order approximation rather than a definitive predictor.\n\n"
    "Next Steps:\n"
    "- Validate the model with additional experimental data.\n"
    "- Explore alternative modeling approaches (e.g., nonlinear regression, machine learning).\n"
    "- Conduct sensitivity analysis to identify the most influential parameters."
)

# Conclusion
doc.add_heading('Conclusion', level=2)
doc.add_paragraph(
    "This report provides a detailed quantitative framework for understanding and optimizing Jatropha oil extraction. While the regression model offers valuable insights, its statistical insignificance and moderate predictive accuracy suggest that further refinement is needed. Nonetheless, the optimization results are promising and highlight the potential for significant yield improvements through process engineering."
)

# Save the document
doc.save('Jatropha_Oil_Extraction_Discussion.docx')
