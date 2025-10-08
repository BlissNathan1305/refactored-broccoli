import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os
import re

# Create output folder
output_folder = "soil_charts"
os.makedirs(output_folder, exist_ok=True)

# Full soil dataset
data = {
    'Location': ['Nung Uyo Idoro', 'Nung Uyo Idoro', 'Ekom Iman', 'Ekom Iman',
                 'Afaha Idoro', 'Afaha Idoro', 'Ikot Idaha', 'Ikot Idaha'],
    'Depth': ['0-50', '50-100'] * 4,
    'Sand (%)': [77.40, 71.44, 77.32, 73.40, 85.36, 79.44, 87.40, 81.32],
    'Silt (%)': [9.64, 15.60, 11.64, 11.52, 5.72, 7.52, 5.52, 7.56],
    'Clay (%)': [12.96, 12.96, 11.04, 15.08, 8.92, 13.04, 7.08, 11.12],
    'TC': ['LS', 'LS', 'LS', 'LS', 'S', 'LS', 'S', 'LS'],
    'pH (H₂O)': [7.08, 7.26, 7.10, 7.37, 7.30, 7.62, 7.82, 5.63],
    'E/C (ds/m)': [0.18, 0.12, 0.13, 0.11, 0.14, 0.17, 0.10, 0.09],
    'OC (%)': [2.19, 1.80, 1.02, 0.96, 2.63, 1.54, 2.23, 0.88],
    'OM (%)': [3.79, 3.11, 1.76, 1.66, 4.55, 2.66, 3.86, 1.52],
    'T/N (%)': [0.09, 0.08, 0.04, 0.04, 0.11, 0.07, 0.10, 0.22],
    'CN:R': [24.3, 22.5, 25.5, 24.0, 23.9, 22.0, 22.3, 22.0],
    'Av.P (mg/kg)': [62.0, 54.0, 56.666, 60.666, 53.999, 43.332, 54.0, 86.666],
    'Ca Cmol/kg': [4.56, 4.08, 4.56, 4.80, 4.32, 6.24, 5.76, 3.34],
    'Mg Cmol/kg': [2.88, 2.06, 2.72, 2.96, 2.68, 4.16, 3.44, 1.94],
    'K Cmol/kg': [0.06, 0.06, 0.03, 0.07, 0.03, 0.08, 0.03, 0.02],
    'Na Cmol/kg': [0.05, 0.03, 0.09, 0.05, 0.09, 0.13, 0.06, 0.07],
    'E/A Cmol/kg': [1.68, 1.20, 2.64, 3.84, 2.64, 4.32, 2.40, 1.68],
    'AL⁺ Cmol/kg': [0.88, 0.66, 1.16, 1.98, 1.12, 3.08, 1.14, 0.92],
    'H⁺ Cmol/kg': [0.80, 0.54, 1.48, 1.86, 1.52, 1.24, 1.26, 0.76],
    'TEB Cmol/kg': [7.55, 6.23, 7.40, 7.88, 7.12, 10.61, 9.29, 5.87],
    'ECEC Cmol/kg': [9.23, 7.43, 10.04, 11.72, 9.76, 14.93, 11.69, 7.55],
    'BSAT (%)': [81.79, 83.84, 73.70, 67.24, 72.95, 71.06, 79.47, 77.74]
}

df = pd.DataFrame(data)

# Sanitize filename to avoid illegal characters
def sanitize_filename(name):
    return re.sub(r'[^\w\-]', '_', name)

# Function to plot and save chart
def plot_and_save(property_name):
    if df[property_name].dtype == 'O':  # Skip non-numeric like 'TC'
        return None
    fig, ax = plt.subplots(figsize=(10, 6))
    for depth in ['0-50', '50-100']:
        subset = df[df['Depth'] == depth]
        ax.bar(subset['Location'], subset[property_name], label=f'Depth {depth}', alpha=0.7)
    
    ax.set_title(f'{property_name} Across Locations and Depths')
    ax.set_ylabel(property_name)
    ax.set_xlabel('Location')
    ax.legend()
    plt.xticks(rotation=45)
    plt.tight_layout()
    
    safe_name = sanitize_filename(property_name)
    filename = f"{output_folder}/{safe_name}.png"
    plt.savefig(filename)
    plt.close()
    return filename

# Create Word document
doc = Document()
doc.add_heading('Comprehensive Soil Physicochemical Properties Report', 0)

# Add charts and captions
for prop in df.columns[2:]:  # Skip Location and Depth
    chart_path = plot_and_save(prop)
    doc.add_heading(prop, level=1)
    if chart_path:
        doc.add_picture(chart_path, width=Inches(6))
        doc.add_paragraph(f"Bar chart showing {prop} variation across locations at depths 0–50 cm and 50–100 cm.")
    else:
        doc.add_paragraph(f"{prop} is a non-numeric parameter and is not visualized as a chart.")

# Save document
doc.save("Full_Soil_Properties_Report.docx")
print("Document exported as 'Full_Soil_Properties_Report.docx'")
