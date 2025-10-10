# export_solar_dryer_fullpaper.py
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# ========== 1. Data (same as before) ==========
times_upper = np.array([0, 60, 120, 180, 240, 300, 360, 420, 480, 540, 600, 660, 720, 780])
mc_upper = np.array([67, 64, 63, 60, 56, 45, 37, 36, 21, 16, 11, 7, 3, 0], dtype=float)

times_lower = np.array([0, 60, 120, 180, 240, 300, 360, 420, 480, 540, 600, 660, 720, 780, 840, 900, 960, 1020, 1140])
mc_lower = np.array([68, 69, 65, 60, 55, 51, 43, 40, 38, 30, 25, 22, 18, 14, 10, 7, 3, 0, 0], dtype=float)

times_open = np.array([0,60,120,180,240,300,360,480,540,600,660,720,780,840,900,960,1020,1080,1140,1200])
mc_open = np.array([52,50,49,47,46,45,44,42,31,26,24,22,19,16,13,9,7,3,3,0], dtype=float)

def compute_drying_rate(time_min, mc_percent):
    dt = np.diff(time_min)
    dmc = np.diff(mc_percent)
    rate = dmc / dt
    t_mid = time_min[:-1] + dt/2
    return t_mid, rate

t_mid_u, rate_u = compute_drying_rate(times_upper, mc_upper)
t_mid_l, rate_l = compute_drying_rate(times_lower, mc_lower)
t_mid_o, rate_o = compute_drying_rate(times_open, mc_open)

# ========== 2. Generate/Save Figures ==========
os.makedirs('figures', exist_ok=True)

# Figure 1: moisture vs time
plt.figure(figsize=(7,4))
plt.plot(times_upper/60, mc_upper, '-o', label='Upper tray')
plt.plot(times_lower/60, mc_lower, '-s', label='Lower tray')
plt.plot(times_open/60, mc_open, '-^', label='Open air')
plt.xlabel('Time (hours)')
plt.ylabel('Moisture content (% wb)')
plt.title('Moisture content vs drying time')
plt.grid(True)
plt.legend()
plt.tight_layout()
plt.savefig('figures/fig1_moisture_vs_time.png', dpi=200)
plt.close()

# Figure 2: drying rate vs time (absolute)
plt.figure(figsize=(7,4))
plt.plot(t_mid_u/60, -rate_u, '-o', label='Upper tray')
plt.plot(t_mid_l/60, -rate_l, '-s', label='Lower tray')
plt.plot(t_mid_o/60, -rate_o, '-^', label='Open air')
plt.xlabel('Time (hours)')
plt.ylabel('Drying rate (%wb per min)')
plt.title('Drying rate vs time')
plt.grid(True)
plt.legend()
plt.tight_layout()
plt.savefig('figures/fig2_drying_rate_vs_time.png', dpi=200)
plt.close()

# ========== 3. Build Document ==========
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
style.font.size = Pt(12)

# Title and authors
h = doc.add_heading('Development and Evaluation of a Direct Passive Polycarbonate Cylindrical Solar Dryer for Crayfish', level=1)
h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p_auth = doc.add_paragraph('Utit, I. I.; Supervisor: Dr. David Onwe', style='Normal')
p_auth.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Abstract
doc.add_heading('Abstract', level=2)
doc.add_paragraph(
    "The need to preserve aquatic proteins such as crayfish in low-electrical settings drives research into passive solar drying. "
    "This study describes the design, construction, and performance evaluation of a direct passive cylindrical solar dryer made from polycarbonate using crayfish as test material. "
    "The dryer attained internal temperatures up to ~54 °C, reduced moisture content from ~48.6% to ~3.51% (upper tray) in ~12 h and from ~42.4% to 2.5% (lower tray) in ~17 h, compared to ~19 h via open-air drying. "
    "Peak drying rates were ~0.081 g H₂O/min and the calculated mass-based efficiency was ~28.3%. "
    "This configuration shows promise for low-cost aquatic product drying in tropical off-grid settings."
)

# Introduction
doc.add_heading('1. Introduction', level=2)
intro_text = (
    "Aquatic protein sources such as crayfish and fish are critical for food security in tropical communities, "
    "but their high moisture content (~70–80 %) makes them highly perishable. Traditional open-sun drying remains widespread, "
    "but suffers contamination, insect infestation, unpredictable weather, and slow drying rates (Younis et al., 2025). "
    "Solar dryers offer improved control, hygiene, and accelerated drying through elevated temperatures and reduced ambient humidity (Fernandes et al., 2024). "
    "Solar drying systems may be classified as direct, indirect, mixed-mode, active or passive (Prakash et al., 2025). "
    "Direct dryers allow product radiation exposure; indirect dryers use heated air streams; mixed-mode combines both; active systems use fans, "
    "while passive rely solely on buoyancy‐driven airflow. Natural convection dryers often exhibit limited airflow under weak thermal gradients (Prakash et al., 2025; Fernandes et al., 2024). "
    "While solar-drying of fruits, grains, and vegetables is widely studied, fewer works address aquatic products in direct passive cylindrical configurations. "
    "Here, we design and evaluate a solar cylinder dryer with polycarbonate glazing and local materials, using crayfish to characterize drying kinetics, thermal performance, and system efficiency."
)
doc.add_paragraph(intro_text)

# Literature Review
doc.add_heading('2. Literature Review', level=2)
lit_text = (
    "Solar drying of fish and aquatic products has been explored in forms such as the direct dryer by Obayopo & Alonge (2018), "
    "which recorded internal temperature increases of ~35 °C and maximum efficiencies exceeding 70% under forced convection. "
    "However, high temperatures risk protein denaturation and lipid oxidation (Fitri et al., 2022). "
    "Hybrid systems combining solar energy with electrical backup have been used to mitigate intermittency (Development of Solar Dryer with Electrical Backup, 2021). "
    "IoT-enabled solar dryers are more recent innovations, allowing real-time monitoring and controlled operation (Modification & Evaluation, 2024). "
    "In numerical modeling, finite-difference greenhouse dryer models (Sadodin & Kashani, 2011) and iterative convective drying estimators (Skarbalius et al., 2021) provide design-driven insight. "
    "Comprehensive reviews place recent advances in glazing, airflow design, hybrid storage, and materials integration at the frontier of solar drying research (Fernandes et al., 2024). "
    "Yet direct passive cylindrical dryers remain comparatively underexplored."
)
doc.add_paragraph(lit_text)

# Materials and Methods
doc.add_heading('3. Materials and Methods', level=2)
mm_text = (
    "The solar dryer constructed is a vertical cylinder using 0.7 mm polycarbonate sheet as the transparent cover, "
    "mounted on a wooden frame. Two black-painted corrugated aluminum absorber plates conform to the cylindrical interior, "
    "with two mesh trays (upper, lower) separated by ~20 cm. Cylinder radius 20.3 cm, height 40.4 cm, giving ~52,297 cm³ volume; glazing area ~0.1294 m²; tray area ~706.95 cm². "
    "Natural passive ventilation (no fan or chimney) was used. Absorbers are painted matte black. The dryer is aligned to optimize sun exposure."
)
doc.add_paragraph(mm_text)

mm2 = (
    "Fresh crayfish (Procambarus spp.) were cleaned and loaded into tray replicates (n=3). Every 60 min, up to equilibrium moisture, "
    "samples from each tray and open-air controls were weighed, and ambient and internal temperatures and humidity logged. "
    "Open-air drying was conducted simultaneously under same interval sampling."
)
doc.add_paragraph(mm2)

mm3 = (
    "Moisture content (wet basis) was computed as MC = (M_w / (M_w + M_d)) ×100. Drying rate was computed as ΔW/Δt (g H₂O per min). "
    "Dryer efficiency (mass-based) was: η = (W L)/(I A t), where W=water mass removed, L=latent heat (≈2,260 kJ/kg), I=solar insolation (W/m²), A=collector area, t=drying duration. "
    "Statistical summaries (mean, standard deviation, min, max) were computed. Moisture and drying-rate curves were plotted."
)
doc.add_paragraph(mm3)

# Results
doc.add_heading('4. Results', level=2)

# Thermal performance
doc.add_heading('4.1 Thermal performance', level=3)
tp_text = (
    "Ambient temperature averaged ~32.77 °C (±6.2 °C). Internal dryer temperature averaged ~40.24 °C (±7.80 °C) at upper tray, "
    "and ~34.83 °C (±17.10 °C) at lower tray. Average temperature differentials (ΔT) were ~16.8 °C (upper) and ~12.6 °C (lower). "
    "Peak internal temperature reached ~54 °C. These results confirm the cylindrical polycarbonate structure and absorbers effectively harnessed solar heat and established a greenhouse microclimate."
)
doc.add_paragraph(tp_text)

# Moisture kinetics
doc.add_heading('4.2 Moisture removal kinetics', level=3)
mk_text = (
    "Upper tray moisture dropped from ~48.6% to ~3.51% in ~12 h. Lower tray dropped from ~42.4% to ~2.50% in ~17 h. Open-air drying reached ~2.77% in ~19 h. "
    "These findings reflect the superior thermal environment of the dryer and the advantage of enclosed drying."
)
doc.add_paragraph(mk_text)

# Embed Figure 1
doc.add_picture('figures/fig1_moisture_vs_time.png', width=Inches(6.0))
doc.add_paragraph('Figure 1. Moisture content vs time for upper tray, lower tray, and open-air.')

# Drying rate
doc.add_heading('4.3 Drying rate behavior', level=3)
dr_text = (
    "Drying-rate curves show an initial relatively high rate (surface moisture removal) followed by a falling-rate regime as internal diffusion becomes limiting. "
    "Peak drying rates: upper tray ~0.081 g/min, lower tray slightly less, open-air ~0.028 g/min. This is consistent with canonical drying theory and confirms that the dryer substantially accelerates moisture removal relative to open-air conditions."
)
doc.add_paragraph(dr_text)

# Embed Figure 2
doc.add_picture('figures/fig2_drying_rate_vs_time.png', width=Inches(6.0))
doc.add_paragraph('Figure 2. Drying rate vs time (absolute values) for upper, lower, and open-air.')

# Efficiency & comparative
doc.add_heading('4.4 Efficiency and comparison', level=3)
cmp_text = (
    "Using total water removal (~42.70 g) and assumed insolation (~188 W/m²), the mass-based dryer efficiency computed is ~28.3%. "
    "Compared with open-air drying, the cylinder dryer reduced drying time by ~7 hours (12 h vs 19 h for equivalent moisture levels). "
    "This underscores the benefit of controlled thermal microclimates and enclosed drying."
)
doc.add_paragraph(cmp_text)

# Discussion
doc.add_heading('5. Discussion', level=2)
dis_text = (
    "The dryer’s ability to elevate internal temperature and maintain moisture gradients illustrates the efficacy of a compact cylindrical geometry with transparent glazing and absorbers. "
    "Efficiency (~28.3%) is moderate, particularly in comparison with forced convection or fan-assisted systems (which may exceed 70%, e.g. Obayopo & Alonge, 2018). "
    "Passive systems must contend with convection limitations under low ΔT. The higher performance of the upper tray points to airflow stratification — additional venting or chimney enhancement may improve uniformity. "
    "The observed drying curves align with two-phase kinetics (constant to falling rate), supporting internal diffusion-limited drying in later stages. "
    "Future improvements could include enhanced ventilation (chimney, vent sizing), testing alternative glazing materials (polycarbonate vs acrylic or glass), collector augmentations, and hybrid designs (fans, thermal storage). "
    "Larger-scale studies, replicate designs, and statistical testing (ANOVA) may bolster robustness and allow generalized predictive models."
)
doc.add_paragraph(dis_text)

# Conclusion
doc.add_heading('6. Conclusion', level=2)
conc_text = (
    "A direct passive cylindrical solar dryer was successfully designed and tested using crayfish as a test commodity. "
    "The system achieved internal temperatures up to ~54 °C, reduced moisture to <4 % within 12–17 h, and demonstrated ~28.3 % mass-based efficiency. "
    "It delivered noteworthy time savings over open-air drying and provided reproducible drying curves. "
    "While the performance is modest relative to active dryers, the low-cost, simple design is promising for off-grid aquatic product drying. "
    "Optimizations of airflow, glazing, and hybrid enhancements may further yield improved efficiencies."
)
doc.add_paragraph(conc_text)

# References heading
doc.add_heading('References', level=2)
refs = [
    "Obayopo, S. O. & Alonge, O. I. (2018). Development and Quality Analysis of a Direct Solar Dryer for Fish. Food and Nutrition Sciences, 9, 474-488. https://doi.org/10.4236/fns.2018.95037",
    "Fitri, N., et al. (2022). A Comprehensive Review on the Processing of Dried Fish. PMC.",
    "Fernandes, L., et al. (2024). A Review on Solar Drying Devices: Heat Transfer, Air … Solar 4(1).",
    "Prakash, R., et al. (2025). A review on natural convective solar dryer. Energy Conversion & Management (forthcoming).",
    "Sadodin, S. & Kashani, T. T. (2011). Numerical investigation of a solar greenhouse tunnel drier for drying of copra. arXiv.",
    "Skarbalius, G., Dziugys, A., Misiulis, E., & Navakas, R. (2021). Iterative method for convective drying of biomass. arXiv."
]
for r in refs:
    doc.add_paragraph(r, style='Normal')

# Save
out = 'solar_dryer_full_paper.docx'
doc.save(out)
print("Document written to:", out)
