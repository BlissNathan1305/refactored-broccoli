# generate_full_book.py
# Full book generator for:
# "Practical Data Analysis with Python and R: A Hands-On Guide"
# Author: Samuel Blessed Nathaniel
# Edition: 2025
#
# Requires: python-docx
# pip install python-docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import date
import textwrap

def set_run_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = name
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_paragraph_with_style(doc, text, style=None, align=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    if align:
        p.alignment = align
    return p

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_code_block(doc, code_text):
    # code block: use a paragraph with Courier New and small shading via prefix and monospace.
    p = doc.add_paragraph()
    run = p.add_run()
    # Pre-format: keep indentation
    run.text = code_text
    run.font.name = 'Courier New'
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9)
    # add a light border by adding empty paragraph before and after
    return p

def add_figure_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Figure placeholder: {}]".format(caption))
    set_run_font(run, name='Times New Roman', size=10, italic=True)
    doc.add_paragraph("Figure: " + caption, style='Caption')

# Build the long content (Parts I - V)
title = "Practical Data Analysis with Python and R: A Hands-On Guide"
author = "Samuel Blessed Nathaniel"
edition = "2025 Edition"
filename = "Practical_Data_Analysis_with_Python_and_R_FULL.docx"

doc = Document()

# ---- Cover Page ----
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover.add_run("\n\n\n")
set_run_font(run, size=18)
run = cover.add_run(title + "\n")
set_run_font(run, size=36, bold=True)
run = cover.add_run("\nA Hands-On Guide\n")
set_run_font(run, size=18, italic=False)
run = cover.add_run("\n\nBy {}\n".format(author))
set_run_font(run, size=18)
run = cover.add_run("\n{}\n".format(edition))
set_run_font(run, size=12)
run = cover.add_run("\n\nSelf-Published\n")
set_run_font(run, size=10, italic=True)
doc.add_page_break()

# ---- Front matter ----
add_heading(doc, "Preface", level=1)
doc.add_paragraph(
    "This book is a practical, hands-on guide to data analysis using Python and R. "
    "It is written for beginners and early-career practitioners who want real examples, "
    "clear explanations, and reproducible code. "
    "You can use the scripts here as a starting point for your projects."
)
doc.add_paragraph(
    "How to use this book: Read sequentially if you're new; otherwise jump to chapters for the tasks you need."
)
doc.add_page_break()

add_heading(doc, "Acknowledgments", level=1)
doc.add_paragraph(
    "Thanks to the open-source community, library authors, and countless contributors. "
    "This guide is a synthesis of best practices, examples, and templates to help you become productive quickly."
)
doc.add_page_break()

# ---- Table of Contents (simple version) ----
add_heading(doc, "Table of Contents", level=1)
toc_items = [
    "Part I – Foundations",
    "  Chapter 1 — Introduction to Data Analysis",
    "  Chapter 2 — Understanding Data",
    "Part II – Working with Data",
    "  Chapter 3 — Data Collection and Importing",
    "  Chapter 4 — Data Cleaning and Preprocessing",
    "  Chapter 5 — Data Exploration (EDA)",
    "Part III – Analysis and Modeling",
    "  Chapter 6 — Statistical Analysis",
    "  Chapter 7 — Data Visualization Masterclass",
    "  Chapter 8 — Predictive Modeling and Machine Learning",
    "Part IV – Reporting, Automation & Exporting Results",
    "  Chapter 9 — Communicating Results",
    "  Chapter 10 — Creating Reports in Python",
    "  Chapter 11 — Creating Reports in R",
    "  Chapter 12 — Building Dashboards",
    "  Chapter 13 — Project Workflow",
    "  Chapter 14 — Best Practices & Ethics",
    "Part V – Advanced Topics",
    "  Chapter 15 — Big Data and Cloud Analysis",
    "  Chapter 16 — AI and Advanced Methods in Analysis",
    "  Chapter 17 — Case Studies",
    "  Chapter 18 — Career Paths in Data Analysis",
    "Back Matter",
    "  Glossary",
    "  References",
    "  Index"
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.style = 'List Number'
doc.add_page_break()

# ---- Content helper to add long paragraphs (wrap) ----
def add_long_text(doc, title_text, paragraphs, level=1):
    add_heading(doc, title_text, level=level)
    for para in paragraphs:
        # wrap long paragraphs for readability within code text
        doc.add_paragraph(para)
    doc.add_page_break()

# -------------------------
# PART I - Foundations
# -------------------------
part_title = "Part I – Foundations"
add_heading(doc, part_title, level=1)

# Chapter 1
chap1_paras = [
    "Data analysis is the process of inspecting, cleaning, transforming, and modeling data with the goal of discovering useful information, informing conclusions, and supporting decision-making. "
    "In practice it means moving from raw inputs (files, tables, or measurements) to decisions or predictions that solve a problem.",
    "Why learn data analysis? Organizations rely on data for decisions: product choices, investment, hiring, healthcare planning, and many other areas. Learning data analysis gives you the ability to convert data into impact.",
    "Why Python and R? Both languages are industry staples. Python is versatile, integrates well with web and production systems, and has strong machine learning frameworks. R is purpose-built for statistical analysis and has a very expressive visualization system. "
]
add_long_text(doc, "Chapter 1 — Introduction to Data Analysis", chap1_paras, level=2)

# Chapter 2
chap2_paras = [
    "Data can be quantitative (numerical) or qualitative (categorical). It can be structured (tables) or unstructured (text, images). Understanding types determines which analysis techniques you should use.",
    "Common data mistakes include wrong data types, missing values, inconsistent category labels, and unhandled outliers. The rest of the book covers practical methods to address these problems.",
    "Quick examples in Python and R are provided throughout to show how to begin exploring and visualizing your data. Practice is essential: follow the examples with your datasets."
]
add_long_text(doc, "Chapter 2 — Understanding Data", chap2_paras, level=2)

# -------------------------
# PART II - Working with Data
# -------------------------
add_heading(doc, "Part II – Working with Data", level=1)

# Chapter 3
chap3_paras = [
    "Data collection is the first active step. Sources include CSV/Excel files, databases, public APIs, web scraping and sensors.",
    "In Python, the pandas library is the workhorse for reading tabular data. In R, readr and readxl do similar tasks. When connecting to databases use SQL connections and fetch only the columns you need."
]
add_long_text(doc, "Chapter 3 — Data Collection and Importing", chap3_paras, level=2)

# Provide code example blocks
doc.add_paragraph("Python example — reading files and APIs:")
py_code = textwrap.dedent("""
import pandas as pd
import requests

# CSV
df_csv = pd.read_csv('sales.csv')

# Excel
df_xlsx = pd.read_excel('survey.xlsx', sheet_name='Sheet1')

# JSON from API
resp = requests.get('https://api.exchangerate-api.com/v4/latest/USD')
data = resp.json()
""")
add_code_block(doc, py_code)

doc.add_paragraph("R example — reading files and APIs:")
r_code = textwrap.dedent("""
library(readr)
library(httr)
df_csv <- read_csv('sales.csv')
df_xlsx <- readxl::read_excel('survey.xlsx')

res <- GET('https://api.exchangerate-api.com/v4/latest/USD')
data <- jsonlite::fromJSON(rawToChar(res$content))
""")
add_code_block(doc, r_code)
doc.add_page_break()

# Chapter 4
chap4_paras = [
    "Cleaning and preprocessing is often the most time-consuming part of any data project. It includes handling missing values, removing duplicate records, converting types, standardizing categories, and dealing with outliers.",
    "Always keep a copy of the raw data and document the cleaning steps. This ensures reproducibility and prevents accidental data loss."
]
add_long_text(doc, "Chapter 4 — Data Cleaning and Preprocessing", chap4_paras, level=2)

doc.add_paragraph("Python example — common cleaning steps:")
py_code2 = textwrap.dedent("""
# Detect missing values
df.isnull().sum()

# Fill numeric missing with mean
df['age'] = df['age'].fillna(df['age'].mean())

# Drop duplicates
df = df.drop_duplicates()

# Convert date column
df['date'] = pd.to_datetime(df['date'])
""")
add_code_block(doc, py_code2)

doc.add_paragraph("R example — common cleaning steps:")
r_code2 = textwrap.dedent("""
# Detect missing values
colSums(is.na(df))

# Fill numeric missing with mean
df$age[is.na(df$age)] <- mean(df$age, na.rm = TRUE)

# Remove duplicates using dplyr
library(dplyr)
df <- distinct(df)

# Convert date
df$date <- as.Date(df$date)
""")
add_code_block(doc, r_code2)
doc.add_page_break()

# Chapter 5
chap5_paras = [
    "Exploratory Data Analysis (EDA) helps you understand patterns, distributions, correlations, and anomalies. Use summary statistics and plots to build intuition before modeling.",
    "Automated tools exist (pandas-profiling / ydata-profiling in Python, DataExplorer in R) and are excellent for a first-pass analysis."
]
add_long_text(doc, "Chapter 5 — Data Exploration (EDA)", chap5_paras, level=2)

doc.add_paragraph("Python example — EDA and visualization:")
py_code3 = textwrap.dedent("""
import matplotlib.pyplot as plt
import seaborn as sns

# Distribution
df['age'].hist()

# Boxplot
sns.boxplot(x=df['salary'])

# Correlation heatmap
sns.heatmap(df.corr(), annot=True)
""")
add_code_block(doc, py_code3)

doc.add_paragraph("R example — EDA and visualization:")
r_code3 = textwrap.dedent("""
library(ggplot2)
# Distribution
ggplot(df, aes(x=age)) + geom_histogram()

# Boxplot
ggplot(df, aes(y=salary)) + geom_boxplot()

# Correlation
library(corrplot)
corrplot(cor(df), method='color')
""")
add_code_block(doc, r_code3)
doc.add_page_break()

# -------------------------
# PART III - Analysis and Modeling
# -------------------------
add_heading(doc, "Part III – Analysis and Modeling", level=1)

# Chapter 6
chap6_paras = [
    "Statistical analysis helps quantify uncertainty and test hypotheses. Key ideas: central tendency (mean/median), spread (variance/SD), confidence intervals, hypothesis tests (t-test, chi-square), correlation, and regression.",
    "Always check assumptions (normality, independence, variance equality) before using parametric tests."
]
add_long_text(doc, "Chapter 6 — Statistical Analysis", chap6_paras, level=2)

doc.add_paragraph("Example — t-test in Python:")
add_code_block(doc, textwrap.dedent("""
from scipy.stats import ttest_ind
group_a = df[df['group']=='A']['score']
group_b = df[df['group']=='B']['score']
t_stat, p_value = ttest_ind(group_a, group_b)
print('p-value:', p_value)
"""))
doc.add_paragraph("Example — t-test in R:")
add_code_block(doc, textwrap.dedent("""
t.test(score ~ group, data=df)
"""))
doc.add_page_break()

# Chapter 7
chap7_paras = [
    "Good visualizations tell a story. Use appropriate chart types and annotate them. Avoid unnecessary decoration, and always label axes and legends.",
    "Use color palettes that work for color-blind readers. Keep the message of each plot clear: what question does this plot answer?"
]
add_long_text(doc, "Chapter 7 — Data Visualization Masterclass", chap7_paras, level=2)

doc.add_paragraph("Python example — bar and line chart:")
add_code_block(doc, textwrap.dedent("""
import matplotlib.pyplot as plt
# Bar chart
df.groupby('category')['value'].mean().plot(kind='bar')
# Line chart
plt.plot(df['date'], df['sales'])
"""))
doc.add_paragraph("R example — ggplot2 examples:")
add_code_block(doc, textwrap.dedent("""
library(ggplot2)
ggplot(df, aes(x=category, y=value)) + geom_col()
ggplot(df, aes(x=date, y=sales)) + geom_line()
"""))
doc.add_page_break()

# Chapter 8
chap8_paras = [
    "Machine learning extends statistical modelling. Supervised learning predicts labels (classification) or quantities (regression). Unsupervised learning finds structure (clustering).",
    "Important steps: prepare data, split into train/test, choose model, evaluate, tune, and save the model for reuse."
]
add_long_text(doc, "Chapter 8 — Predictive Modeling and Machine Learning", chap8_paras, level=2)

doc.add_paragraph("Python example — simple pipeline with scikit-learn:")
add_code_block(doc, textwrap.dedent("""
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score

X = df[['feature1','feature2']]
y = df['target']
X_train, X_test, y_train, y_test = train_test_split(X,y,test_size=0.2, random_state=42)

model = LinearRegression()
model.fit(X_train, y_train)
pred = model.predict(X_test)
print('R2:', r2_score(y_test, pred))
"""))

doc.add_paragraph("R example — caret example:")
add_code_block(doc, textwrap.dedent("""
library(caret)
set.seed(42)
train_index <- createDataPartition(df$target, p=0.8, list=FALSE)
train <- df[train_index,]
test <- df[-train_index,]
model <- train(target ~ ., data=train, method='lm')
pred <- predict(model, test)
postResample(pred, test$target)
"""))
doc.add_page_break()

# -------------------------
# PART IV - Reporting & Automation
# -------------------------
add_heading(doc, "Part IV – Reporting, Automation & Exporting Results", level=1)

# Chapter 9
chap9_paras = [
    "Reporting converts analysis into action. Executive summaries highlight key findings concisely. Include visuals, short explanations, and explicit recommendations.",
    "Produce reproducible reports; attach code and data descriptions in an appendix."
]
add_long_text(doc, "Chapter 9 — Communicating Results Effectively", chap9_paras, level=2)

# Chapter 10
chap10_paras = [
    "In Python, python-docx can produce Word documents. ReportLab or wkhtmltopdf produce PDFs. Jupyter notebooks exported to HTML are useful for interactive sharing. Streamlit builds web apps for dashboards.",
    "Automate reporting with schedulers, cron, or cloud functions so stakeholders get regular updates."
]
add_long_text(doc, "Chapter 10 — Creating Reports in Python", chap10_paras, level=2)

doc.add_paragraph("Python example — python-docx snippet:")
add_code_block(doc, textwrap.dedent("""
from docx import Document
doc = Document()
doc.add_heading('Sales Report', 0)
doc.add_paragraph('Executive Summary: ...')
doc.save('report.docx')
"""))

# Chapter 11
chap11_paras = [
    "R Markdown is the best-in-class tool for producing reproducible reports and can render to Word, PDF, or HTML. The officer package writes DOCX programmatically for templated reports."
]
add_long_text(doc, "Chapter 11 — Creating Reports in R", chap11_paras, level=2)

doc.add_paragraph("R example — minimal R Markdown header:")
add_code_block(doc, textwrap.dedent("""
---
title: "Sales Report"
output: word_document
---
# Analysis
```{r}
library(ggplot2)
