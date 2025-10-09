from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()
doc.add_heading('Comprehensive Discussion of Soil Physicochemical Properties', 0)

# Define a helper function to add paragraphs with formatting
def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold

# Introduction
add_paragraph("This document presents a detailed discussion of the soil physicochemical properties measured across various locations and depths (0–50 cm and 50–100 cm). These properties are essential for understanding soil fertility, structure, and suitability for agricultural and environmental applications.", bold=True)

# Texture Properties
add_paragraph("\n1. Soil Texture Components (Sand, Silt, Clay):", bold=True)
add_paragraph("Soil texture is determined by the relative proportions of sand, silt, and clay. These components influence water retention, drainage, aeration, and nutrient availability.")
add_paragraph("- Sand (%): High sand content indicates better drainage but lower nutrient and water retention. Variation across depths may reflect parent material or erosion.")
add_paragraph("- Silt (%): Silt contributes to soil fertility and water retention. Changes across depths can signal deposition or weathering processes.")
add_paragraph("- Clay (%): Clay-rich soils retain nutrients and water effectively but may suffer from poor aeration. Depth variation may indicate pedogenic processes.")

# Chemical Properties
add_paragraph("\n2. pH (H₂O):", bold=True)
add_paragraph("Soil pH affects nutrient availability and microbial activity. Values below 6.5 suggest acidic conditions, which may limit phosphorus and molybdenum availability. Depth variation may be due to leaching or organic matter decomposition.")

add_paragraph("\n3. Electrical Conductivity (E/C ds/m):", bold=True)
add_paragraph("E/C measures soil salinity. High values can impair plant growth. Variation across depths may indicate salt accumulation or irrigation practices.")

add_paragraph("\n4. Organic Carbon (OC %) and Organic Matter (OM %):", bold=True)
add_paragraph("OC and OM are indicators of soil fertility and biological activity. Higher values in topsoil (0–50 cm) are expected due to organic inputs. Decline with depth reflects reduced biological activity.")

add_paragraph("\n5. Total Nitrogen (T/N %):", bold=True)
add_paragraph("T/N is crucial for plant growth. Its distribution mirrors organic matter content. Low nitrogen may require fertilization for optimal crop yield.")

add_paragraph("\n6. Carbon to Nitrogen Ratio (CN:R):", bold=True)
add_paragraph("CN ratio affects decomposition rates and nitrogen mineralization. Ratios above 20 may slow decomposition, while lower ratios promote microbial activity.")

add_paragraph("\n7. Available Phosphorus (Av.P mg/kg):", bold=True)
add_paragraph("Phosphorus is vital for root development. Its availability is pH-dependent. Low values may necessitate phosphate fertilization.")

# Exchangeable Cations
add_paragraph("\n8. Exchangeable Cations (Ca, Mg, K, Na Cmol/kg):", bold=True)
add_paragraph("These cations influence soil structure and fertility.")
add_paragraph("- Calcium (Ca): Essential for root development and soil aggregation.")
add_paragraph("- Magnesium (Mg): Important for photosynthesis and enzyme activation.")
add_paragraph("- Potassium (K): Regulates water uptake and enzyme activity.")
add_paragraph("- Sodium (Na): Excessive Na can lead to poor soil structure and salinity issues.")

add_paragraph("\n9. Exchange Acidity (E/A), Aluminum (AL⁺), and Hydrogen (H⁺) Cmol/kg:", bold=True)
add_paragraph("These parameters reflect soil acidity. High AL⁺ and H⁺ levels can be toxic to plants and reduce nutrient availability.")

# Base Saturation and Capacity
add_paragraph("\n10. Total Exchangeable Bases (TEB) and Effective Cation Exchange Capacity (ECEC) Cmol/kg:", bold=True)
add_paragraph("TEB and ECEC indicate the soil’s ability to retain and supply nutrients. Higher values suggest better fertility and buffering capacity.")

add_paragraph("\n11. Base Saturation (BSAT %):", bold=True)
add_paragraph("BSAT reflects the proportion of exchange sites occupied by base cations. High BSAT is desirable for nutrient-rich soils.")

# TC (Non-numeric)
add_paragraph("\n12. TC (Total Carbon):", bold=True)
add_paragraph("TC is a qualitative parameter that complements OC and OM. It reflects the total carbon pool, including inorganic forms.")

# Conclusion
add_paragraph("\nConclusion:", bold=True)
add_paragraph("The soil properties analyzed provide a comprehensive understanding of soil health and fertility across different locations and depths. These insights are crucial for land management, crop selection, and sustainable agricultural practices.")

# Save the document
doc.save("Soil_Properties_Discussion.docx")
print("Document saved as 'Soil_Properties_Discussion.docx'")
